"""
Generate the full Brain Tumor Detection project report as a Word .docx file.

Usage:
    cd "C:\\Brain Tumor Detection"
    .\.venv\Scripts\python.exe docs\generate_report.py

Output:
    docs/PROJECT_REPORT.docx
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).resolve().parent / "PROJECT_REPORT.docx"

# ------------------------------------------------------------------ helpers
def _set_cell(cell, text, bold=False, align="left"):
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.bold = bold
    run.font.size = Pt(10)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _set_cell(t.rows[0].cells[i], h, bold=True, align="center")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _set_cell(t.rows[ri + 1].cells[ci], str(val))
    doc.add_paragraph()


def heading(doc, level, text):
    doc.add_heading(text, level=level)

def para(doc, text, bold=False, italic=False, align="left"):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(11)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)

def spaced(doc):
    doc.add_paragraph()

# ------------------------------------------------------------------ build
def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for i in range(1, 5):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)

    # ============================================================ TITLE PAGE
    for _ in range(6):
        doc.add_paragraph()
    para(doc,
         "Brain Tumor Detection from MRI Images using an "
         "Ensemble of CNN, Transfer Learning and "
         "Swin Transformer Models",
         bold=True, align="center")
    spaced(doc)
    para(doc, "8th Semester Project Report", align="center")
    para(doc, "submitted to", italic=True, align="center")
    para(doc, "<UNIVERSITY_NAME>", bold=True, align="center")
    spaced(doc)
    para(doc,
         "In partial fulfilment of the requirements for the award of the degree of",
         align="center")
    para(doc, "Bachelor of Technology", bold=True, align="center")
    para(doc, "in", italic=True, align="center")
    para(doc, "<SPECIALIZATION>", bold=True, align="center")
    spaced(doc)
    para(doc, "by", italic=True, align="center")
    para(doc, "<STUDENT_NAME_1> (Roll No. <ROLL_1>)", align="center")
    para(doc, "<STUDENT_NAME_2> (Roll No. <ROLL_2>)", align="center")
    para(doc, "<STUDENT_NAME_3> (Roll No. <ROLL_3>)", align="center")
    para(doc, "<STUDENT_NAME_4> (Roll No. <ROLL_4>)", align="center")
    spaced(doc)
    para(doc, "Under the guidance of", italic=True, align="center")
    para(doc, "<SUPERVISOR_NAME>, <SUPERVISOR_DESIGNATION>", bold=True, align="center")
    spaced(doc)
    para(doc, "Department of <DEPARTMENT_NAME>", align="center")
    para(doc, "<UNIVERSITY_NAME>", bold=True, align="center")
    para(doc, "<SUBMISSION_MONTH_YEAR>", align="center")

    doc.add_page_break()

    # ============================================================ DECLARATION
    heading(doc, 1, "Declaration of Authorship")
    para(doc,
         'We hereby declare that the project report entitled '
         '"Brain Tumor Detection from MRI Images using an Ensemble of CNN, '
         'Transfer Learning and Swin Transformer Models" is an authentic '
         'record of our own work carried out at the Department of '
         '<DEPARTMENT_NAME>, <UNIVERSITY_NAME>, during the 8th semester of '
         'the academic year <ACADEMIC_YEAR> under the supervision of '
         '<SUPERVISOR_NAME>, <SUPERVISOR_DESIGNATION>.',
         align="justify")
    para(doc,
         "We further declare that the matter embodied in this project report "
         "has not been submitted by us for the award of any other degree or "
         "diploma of this or any other Institute or University. All the "
         "information has been obtained and presented in accordance with "
         "academic rules and ethical conduct. We also declare that, as "
         "required by these rules and conduct, we have fully cited and "
         "referenced all materials and results that are not original to "
         "this work. The dataset used in this project is publicly available "
         "under its original license; no patient-identifiable data was "
         "collected, generated, or stored.",
         align="justify")
    spaced(doc)
    para(doc, "Place: <UNIVERSITY_ADDRESS>")
    para(doc, "Date: <SUBMISSION_MONTH_YEAR>")
    spaced(doc)
    para(doc, "Signatures:", bold=True)
    spaced(doc)
    para(doc, "___________________          ___________________")
    para(doc, "<STUDENT_NAME_1>              <STUDENT_NAME_2>")
    para(doc, "Roll No: <ROLL_1>              Roll No: <ROLL_2>")
    spaced(doc)
    para(doc, "___________________          ___________________")
    para(doc, "<STUDENT_NAME_3>              <STUDENT_NAME_4>")
    para(doc, "Roll No: <ROLL_3>              Roll No: <ROLL_4>")

    doc.add_page_break()

    # ============================================================ CERTIFICATE
    heading(doc, 1, "Certificate of Recommendation")
    para(doc,
         'This is to certify that the Dissertation Report entitled '
         '"Brain Tumor Detection from MRI Images using an Ensemble of CNN, '
         'Transfer Learning and Swin Transformer Models" submitted by '
         '<STUDENT_NAME_1>, <STUDENT_NAME_2>, <STUDENT_NAME_3>, and '
         '<STUDENT_NAME_4> to <UNIVERSITY_NAME> is a record of bonafide '
         'project work carried out by them under my supervision and guidance, '
         'and is worthy of consideration for the award of the degree of '
         'Bachelor of Technology (B.Tech) in <SPECIALIZATION>.',
         align="justify")
    spaced(doc)
    para(doc, "<SUPERVISOR_NAME>,", bold=True)
    para(doc, "Project Supervisor,")
    para(doc, "<SUPERVISOR_DESIGNATION>, Department of <DEPARTMENT_NAME>,")
    para(doc, "<UNIVERSITY_NAME>.")
    spaced(doc)
    para(doc, "Approved By:", bold=True)
    spaced(doc)
    para(doc, "<HOD_NAME>,", bold=True)
    para(doc, "HoD / Reviewer, Department of <DEPARTMENT_NAME>,")
    para(doc, "<UNIVERSITY_NAME>.")

    doc.add_page_break()

    # ============================================================ ACKNOWLEDGEMENT
    heading(doc, 1, "Acknowledgement")
    para(doc,
         "We would like to first express our sincere gratitude to our project "
         "supervisor, <SUPERVISOR_NAME>, <SUPERVISOR_DESIGNATION>, Department "
         "of <DEPARTMENT_NAME>, <UNIVERSITY_NAME>. Their continued guidance, "
         "technical insight, and constructive feedback during every stage of "
         "this project were instrumental in shaping the final outcome. Their "
         "willingness to discuss difficult engineering decisions, including "
         "the diagnosis of subtle numerical-stability issues during "
         "mixed-precision training, helped us approach the work with both "
         "rigour and curiosity.",
         align="justify")
    para(doc,
         "We take this opportunity to express our gratitude to all faculty "
         "members of the Department of <DEPARTMENT_NAME> for their support, "
         "the encouragement they provided during our coursework, and the "
         "foundation they laid in machine learning, computer vision, and "
         "software engineering.",
         align="justify")
    para(doc,
         "We acknowledge the open-source community whose tools made this "
         "work feasible: the PyTorch and timm library maintainers, the "
         "authors of pytorch-grad-cam, the Streamlit and FastAPI teams, "
         "and Masoud Nickparvar for releasing the Brain Tumor MRI Dataset "
         "on Kaggle under a permissive license that supports academic "
         "research.",
         align="justify")
    para(doc,
         "Finally, we thank our parents and families for their unceasing "
         "encouragement, patience, and support throughout the semester.",
         align="justify")

    doc.add_page_break()

    # ============================================================ ABSTRACT
    heading(doc, 1, "Abstract")
    para(doc,
         "The accurate and timely classification of brain tumors from "
         "magnetic resonance imaging (MRI) is a clinically important problem, "
         "yet manual interpretation by radiologists is time-consuming, "
         "expertise-intensive, and subject to inter-observer variability. "
         "Reported inter-rater agreement on tumor type lies in the 70 to 85 "
         "percent range, while access to expert neuroradiologists remains "
         "uneven across geographies. There is a clear case for AI-assisted "
         "decision support that is fast, consistent, and explainable.",
         align="justify")
    para(doc,
         "This project presents an end-to-end deep-learning system that "
         "classifies T1-weighted brain MRI slices into four clinical "
         "categories: glioma, meningioma, pituitary tumor, and no tumor. "
         "The system uses an ensemble of three architecturally diverse "
         "models. The first model is a custom convolutional neural network "
         "(CNN) of approximately 1.2 million parameters, trained from "
         "random initialisation. The second is EfficientNet-B0, adapted via "
         "two-stage transfer learning from ImageNet, totalling about 4.0 "
         "million parameters. The third is Swin-Tiny, a hierarchical vision "
         "transformer of 27.5 million parameters that uses windowed "
         "self-attention and is similarly fine-tuned in two stages. "
         "The three models' softmax probability outputs are combined by a "
         "weighted-average ensemble whose mixing weights were obtained by "
         "exhaustive grid search on a held-out validation set.",
         align="justify")
    para(doc,
         "The system was developed and evaluated on the publicly available "
         "Brain Tumor MRI Dataset (Nickparvar, Kaggle, 7,200 images across "
         "four classes). A stratified 80/20 split partitions the 5,600 "
         "training images into 4,480 training and 1,120 validation samples, "
         "while 1,600 testing images remain held out and are evaluated "
         "exactly once. On this held-out test set, the strongest individual "
         "model (EfficientNet-B0) achieves 94.94% accuracy and a "
         "macro-averaged ROC-AUC of 0.991; the three-model ensemble achieves "
         "94.69% accuracy at 0.991 AUC. We report an honest empirical "
         "finding: that ensembling matches but does not exceed the dominant "
         "base learner, and we discuss it through the lens of the "
         "dominant-model problem in ensemble learning.",
         align="justify")
    para(doc,
         "Beyond accuracy, the system provides per-prediction Grad-CAM "
         "heatmaps for all three models, a low-confidence threshold gate "
         "that explicitly recommends radiologist review for uncertain cases, "
         "and four deployment channels: a Streamlit web interface, a FastAPI "
         "REST backend, ONNX exports that yield a 2 to 3.5x CPU speedup, "
         "and a Docker container.",
         align="justify")
    spaced(doc)
    para(doc,
         "Keywords: Brain Tumor Classification, Magnetic Resonance Imaging, "
         "Convolutional Neural Networks, Transfer Learning, "
         "Swin Transformer, Ensemble Learning, Grad-CAM, "
         "Computer-Aided Diagnosis.",
         bold=True)

    doc.add_page_break()

    # ============================================================ CHAPTER 1
    heading(doc, 1, "Chapter 1: Introduction")

    heading(doc, 2, "1.1 Medical Imaging and Computer-Aided Diagnosis")
    para(doc,
         "Brain tumors are abnormal growths of cells that form inside the "
         "cranium. They can originate from brain tissue itself (primary "
         "tumors) or spread to the brain from other parts of the body "
         "(secondary or metastatic tumors). The World Health Organization "
         "(WHO) classifies primary brain tumors into over 150 subtypes "
         "based on cell of origin, molecular markers, and grade of "
         "malignancy. Among these, gliomas, meningiomas, and pituitary "
         "adenomas are the three most commonly encountered categories in "
         "clinical practice.",
         align="justify")
    para(doc,
         "Magnetic Resonance Imaging (MRI) is the preferred diagnostic "
         "modality for brain tumors because it provides excellent "
         "soft-tissue contrast without exposing the patient to ionising "
         "radiation. A standard brain MRI examination produces dozens of "
         "two-dimensional slices across multiple sequences (T1-weighted, "
         "T2-weighted, FLAIR, and contrast-enhanced T1), each of which "
         "must be visually inspected by a neuroradiologist. Manual "
         "interpretation of a complete scan typically takes 15 to 30 "
         "minutes per patient, and reported inter-observer agreement on "
         "tumor type ranges from 70 to 85 percent, meaning that even "
         "experienced radiologists disagree in a non-trivial fraction of "
         "cases [1].",
         align="justify")
    para(doc,
         "Computer-Aided Diagnosis (CAD) systems aim to assist radiologists "
         "by providing automated second-reader opinions. Over the past "
         "decade, deep learning, and convolutional neural networks (CNNs) in "
         "particular, have transformed the CAD landscape. A well-trained "
         "deep learning model can analyse an MRI slice in milliseconds, "
         "provide a confidence-calibrated prediction, and highlight the "
         "regions of the image that most influenced its decision. These "
         "properties make deep-learning-based CAD systems a promising tool "
         "for triaging scans in high-volume radiology departments and for "
         "providing diagnostic support in resource-limited settings where "
         "specialist radiologists are scarce [2].",
         align="justify")
    para(doc,
         "India records approximately 28,000 new brain tumor cases per year "
         "according to the International Association of Cancer Registries "
         "and the Indian Journal of Neurology (2025), with some estimates "
         "reaching 40,000 to 50,000 cases annually. Brain and central "
         "nervous system tumors account for roughly 2 to 3 percent of all "
         "cancers in the country. Meanwhile, India has fewer than one "
         "radiologist per 100,000 people, with the shortage being far more "
         "severe in rural and semi-urban areas [17]. This imbalance between "
         "growing diagnostic demand and limited specialist supply makes the "
         "development of reliable AI-assisted tools not merely an academic "
         "exercise but a public health priority.",
         align="justify")

    heading(doc, 2, "1.2 Objectives of the Project")
    heading(doc, 3, "1.2.1 Primary Objective")
    para(doc,
         "To develop a deep-learning-based brain tumor classification system "
         "that takes a single T1-weighted brain MRI slice as input and "
         "predicts one of four diagnostic categories (glioma, meningioma, "
         "pituitary tumor, or no tumor) with high accuracy, calibrated "
         "confidence, and visual explainability, using an ensemble of three "
         "architecturally diverse neural networks.",
         align="justify")

    heading(doc, 3, "1.2.2 Secondary Objectives")
    bullet(doc,
           "Objective 1: To design and train a custom CNN from scratch to "
           "serve as a baseline model and to provide architectural diversity "
           "for the ensemble.")
    bullet(doc,
           "Objective 2: To adapt a pretrained EfficientNet-B0 model for "
           "brain tumor classification using a two-stage transfer learning "
           "strategy (freeze-then-fine-tune) that maximises accuracy while "
           "avoiding catastrophic forgetting of pretrained features.")
    bullet(doc,
           "Objective 3: To fine-tune a Swin-Tiny vision transformer for "
           "the same classification task, introducing self-attention-based "
           "architectural diversity alongside the CNN-based models.")
    bullet(doc,
           "Objective 4: To combine the three models into a weighted-average "
           "ensemble whose weights are determined by exhaustive grid search "
           "on the validation set, and to compare this ensemble against soft "
           "voting and stacking baselines.")
    bullet(doc,
           "Objective 5: To implement Gradient-weighted Class Activation "
           "Mapping (Grad-CAM) for all three models to provide visual "
           "explanations that highlight which regions of the MRI the model "
           "attended to when making its prediction.")
    bullet(doc,
           "Objective 6: To build a production-ready Streamlit web interface "
           "and a FastAPI REST backend that wrap the trained models and "
           "expose real-time prediction with Grad-CAM overlays.")
    bullet(doc,
           "Objective 7: To export all three models to ONNX format for "
           "cross-platform deployment and to verify that ONNX inference "
           "preserves classification accuracy while delivering faster "
           "CPU-based inference.")
    bullet(doc,
           "Objective 8: To rigorously evaluate the system on a held-out "
           "test set of 1,600 images that is never used during training or "
           "hyperparameter tuning, and to report honest metrics including "
           "accuracy, macro F1-score, macro ROC-AUC, per-class confusion "
           "matrices, and inference latency.")

    heading(doc, 2, "1.3 Challenges in Brain Tumor Classification")
    para(doc,
         "Automated brain tumor classification from MRI faces several "
         "technical and practical challenges that any robust system must "
         "address:",
         align="justify")
    bullet(doc,
           "High intra-class variability: Tumors of the same type can "
           "appear very different across patients depending on size, "
           "location, growth stage, and scanner settings. A glioma in one "
           "patient may look visually distinct from a glioma in another.")
    bullet(doc,
           "Inter-class visual similarity: Certain tumor types, "
           "particularly gliomas and meningiomas, can appear remarkably "
           "similar on a single T1-weighted slice. Even experienced "
           "radiologists sometimes require additional sequences (T2, FLAIR, "
           "contrast-enhanced) to distinguish them.")
    bullet(doc,
           "Limited training data: Medical imaging datasets are inherently "
           "small compared to the millions of images available for natural "
           "image tasks like ImageNet. With approximately 5,600 training "
           "images, overfitting is a real risk that must be mitigated "
           "through augmentation, regularisation, and transfer learning.")
    bullet(doc,
           "Scanner and protocol heterogeneity: Images acquired on "
           "different MRI scanners, with different field strengths, "
           "different coils, and different acquisition protocols produce "
           "systematically different intensity profiles. A model trained on "
           "images from one scanner may degrade in performance when applied "
           "to images from another.")
    bullet(doc,
           "Explainability and clinical trust: Clinicians are unlikely to "
           "adopt a diagnostic tool that functions as a black box. For a "
           "brain tumor classifier to be clinically useful, it must provide "
           "not only a prediction but also a human-interpretable explanation "
           "of why that prediction was made. Grad-CAM heatmaps address this "
           "requirement by showing which image regions drove the decision.")
    bullet(doc,
           "Calibrated uncertainty: A system that outputs overconfident "
           "incorrect predictions is worse than one that acknowledges "
           "uncertainty. The system must be able to flag low-confidence "
           "predictions and recommend radiologist review rather than "
           "presenting every output with equal conviction.")

    heading(doc, 2, "1.4 Methodologies")
    para(doc,
         "The methodology of this project follows a structured, multi-stage "
         "pipeline that proceeds from data acquisition and preparation "
         "through model training, ensemble construction, explainability, "
         "and deployment. The stages are summarised below; each is described "
         "in full detail in the relevant subsequent chapter.",
         align="justify")
    bullet(doc,
           "Data Acquisition: The Brain Tumor MRI Dataset (Nickparvar, "
           "Kaggle) is used, containing 7,200 labelled MRI images across "
           "four classes. The dataset is split into training (4,480 images), "
           "validation (1,120 images), and test (1,600 images) subsets using "
           "stratified random sampling with a fixed seed for "
           "reproducibility. Full details are given in Chapter 3.")
    bullet(doc,
           "Preprocessing and Augmentation: All images are converted to "
           "RGB, resized to 224 by 224 pixels, and normalised using ImageNet "
           "statistics. Training images are additionally augmented with "
           "random horizontal flips, small rotations, affine "
           "transformations, brightness and contrast jitter, and random "
           "erasing to improve generalisation. Details are in Chapter 3.")
    bullet(doc,
           "Model Training: Three architecturally diverse models are "
           "trained: a custom 4-block CNN (1.2 million parameters), an "
           "EfficientNet-B0 with two-stage transfer learning (4.0 million "
           "parameters), and a Swin-Tiny transformer with two-stage "
           "fine-tuning (27.5 million parameters). All three use the same "
           "AdamW optimiser, CosineAnnealingLR scheduler, bfloat16 mixed "
           "precision, and early stopping. Details are in Chapter 5.")
    bullet(doc,
           "Ensemble Construction: Each trained model produces a "
           "four-element softmax probability vector. Three ensemble "
           "combination methods are evaluated: soft voting, weighted "
           "average, and stacking. Optimal weights for the weighted "
           "average are found by exhaustive grid search over the simplex "
           "at step 0.05 (231 candidates). Details are in Chapter 5.")
    bullet(doc,
           "Explainability: Grad-CAM heatmaps are generated for each base "
           "model to visualise which regions of the MRI contributed most "
           "to the predicted class. These overlays are surfaced in both the "
           "Streamlit UI and the FastAPI response. Details are in Chapter 5.")
    bullet(doc,
           "Deployment: The trained models are wrapped in four deployment "
           "interfaces: a Streamlit web app for interactive demos, a FastAPI "
           "REST service for programmatic access, ONNX exports for "
           "cross-platform inference, and a Docker container for cloud "
           "deployment. Details are in Chapters 4 and 5.")

    doc.add_page_break()

    # ============================================================ CHAPTER 2
    heading(doc, 1, "Chapter 2: Literature Survey")

    heading(doc, 2, "2.1 Overview of the Field")
    para(doc,
         "Two recent systematic reviews provide a useful starting point. "
         "Bouhafra and El Bahi (2025) reviewed 60 papers on deep-learning "
         "brain tumor classification published between 2020 and January "
         "2024, covering transfer learning, autoencoders, transformers, "
         "and attention mechanisms [1]. Ranjbarzadeh et al. (2025) "
         "published a comprehensive survey of explainable AI and "
         "vision-transformer approaches specifically for brain tumor "
         "detection [2]. These reviews establish three dominant research "
         "themes: (1) transfer learning from ImageNet-pretrained "
         "backbones (predominantly EfficientNet and ResNet families), "
         "(2) vision-transformer architectures (especially the Swin "
         "family), and (3) ensemble methods that combine multiple "
         "architectures.",
         align="justify")
    para(doc,
         "Section 2.2 reviews ten of the most directly relevant papers "
         "in detail. Each subsection presents the paper's methodology, "
         "dataset, headline results, and the specific way in which the "
         "paper informed the design of our own project. Section 2.3 "
         "consolidates the reviewed work in a single comparison table.",
         align="justify")

    heading(doc, 2, "2.2 Review of Related Work")

    heading(doc, 3, "2.2.1 Babu Vimala et al. (2023): Hybrid Deep Learning "
            "with EfficientNet Variants")
    para(doc,
         "Babu Vimala et al. proposed a transfer-learning approach using "
         "fine-tuned EfficientNet variants (B0 through B4) on the CE-MRI "
         "Figshare dataset for three-class brain tumor classification "
         "(glioma, meningioma, pituitary). Their methodology employed a "
         "two-step refinement: initialising models with ImageNet "
         "weights, then adding custom top layers for tumor "
         "classification. Grad-CAM visualisation was incorporated to "
         "highlight attention maps. The best-performing model, "
         "EfficientNetB2, achieved 99.06% test accuracy with 98.79% "
         "F1-score. This study established EfficientNet as a strong "
         "baseline for medical image classification when combined with "
         "appropriate fine-tuning and data augmentation. Our project "
         "draws on this work in its use of two-stage transfer learning "
         "for the EfficientNet-B0 base model [3].",
         align="justify")

    heading(doc, 3, "2.2.2 Islam et al. (2024): BrainNet Optimised "
            "EfficientNet Architecture")
    para(doc,
         "Islam et al. presented BrainNet, an optimised EfficientNet "
         "pipeline evaluated on the CE-MRI dataset of 3,064 T1-weighted "
         "images. The study compared all EfficientNet variants from B0 "
         "to B7 and incorporated advanced preprocessing and "
         "augmentation techniques. The best model, EfficientNetB3, "
         "achieved 99.69% test accuracy. The paper emphasises the role "
         "of learning-rate scheduling and progressive unfreezing in "
         "achieving top performance on small medical imaging datasets. "
         "This work directly informed our project's two-stage training "
         "schedule, where Stage 1 trains only the head and Stage 2 "
         "unfreezes the backbone with a 10x smaller learning rate [4].",
         align="justify")

    heading(doc, 3, "2.2.3 Alnowami et al. (2024): EfficientNetv2 with "
            "Attention Mechanisms")
    para(doc,
         "Alnowami et al. enhanced the EfficientNetv2 architecture with "
         "a Global Attention Mechanism (GAM) and Efficient Channel "
         "Attention (ECA) for MRI-based brain tumor classification. The "
         "integration of attention mechanisms allowed the model to "
         "focus on salient features within complex MRI images, "
         "improving classification accuracy on a four-class dataset. "
         "The model achieved 99.76% accuracy, setting a new benchmark "
         "at the time of publication. The work also incorporated "
         "Grad-CAM visualisation for clinical interpretability. This "
         "study demonstrates that attention-augmented CNNs can rival "
         "full transformer architectures at lower computational cost, "
         "an insight relevant to our choice of EfficientNet-B0 over "
         "heavier transformer models given the 4 GB VRAM constraint [5].",
         align="justify")

    heading(doc, 3, "2.2.4 Liu et al. (2021): Original Swin Transformer")
    para(doc,
         "Liu et al. introduced the Swin Transformer, a hierarchical "
         "vision transformer that uses shifted windows for efficient "
         "self-attention. Unlike earlier vision transformers that "
         "compute attention over the entire image (quadratic cost), "
         "Swin partitions the image into non-overlapping windows and "
         "computes attention only within each window. In alternating "
         "layers, the windows are shifted by half their size, allowing "
         "information to flow across window boundaries. This design "
         "yields linear computational complexity in image size while "
         "preserving global modelling capability. The Swin-Tiny variant "
         "achieves 81.3% top-1 accuracy on ImageNet at only 28 million "
         "parameters. This paper is the foundation for our third base "
         "model [15].",
         align="justify")

    heading(doc, 3, "2.2.5 Haq et al. (2024): Modified Swin Transformer "
            "with HSW-MSA and ResMLP")
    para(doc,
         "Haq et al. proposed a modified Swin Transformer for brain "
         "tumor diagnosis, introducing a Hybrid Shifted Windows "
         "Multi-Head Self-Attention (HSW-MSA) module along with a "
         "Residual MLP (ResMLP) replacing the traditional MLP layer. "
         "These modifications aimed to improve classification accuracy, "
         "reduce memory usage, and simplify training complexity. The "
         "proposed model was evaluated on a publicly available "
         "four-class brain MRI dataset using transfer learning and "
         "data augmentation. The model achieved a remarkable 99.92% "
         "accuracy, surpassing previous CNN-based methods. This paper "
         "demonstrates the strong potential of transformer architectures "
         "for medical image classification and motivated our inclusion "
         "of Swin-Tiny as one of the three base models in our "
         "ensemble [6].",
         align="justify")

    heading(doc, 3, "2.2.6 Alsubai et al. (2024): Hybrid Swin Transformer "
            "and ResNet50V2")
    para(doc,
         "Alsubai et al. proposed a hybrid architecture combining the "
         "Swin Transformer with ResNet50V2 for enhanced brain tumor "
         "classification from MRI images. The Swin Transformer captures "
         "local-and-global hierarchical features through windowed "
         "self-attention, while ResNet50V2 provides complementary deep "
         "convolutional features. The two streams are fused through "
         "transfer learning to leverage pretrained weights from both "
         "backbones. The hybrid architecture demonstrates that "
         "combining local feature extraction (CNN) with global context "
         "modelling (transformer) yields robust classification across "
         "diverse MRI acquisition conditions. This dual-stream "
         "philosophy parallels our project's ensemble approach, which "
         "combines three architecturally distinct models [7].",
         align="justify")

    heading(doc, 3, "2.2.7 Khan et al. (2023): Weighted Average Ensemble "
            "Deep Learning")
    para(doc,
         "Khan et al. proposed a weighted average ensemble model "
         "combining three deep-learning feature spaces: a VGG19 model "
         "(transfer learning), a standard CNN, and a CNN with data "
         "augmentation. The optimal combination of weights was found "
         "by exhaustive grid search. The dataset used was The Cancer "
         "Genome Atlas (TCGA) lower-grade glioma collection with 3,929 "
         "MRI images. The ensemble outperformed each individual model "
         "in accuracy, precision, and F1-score, demonstrating that "
         "combining models with different feature spaces can compensate "
         "for individual model weaknesses. This paper directly "
         "motivates our project's grid-search approach to "
         "ensemble-weight optimisation [8].",
         align="justify")

    heading(doc, 3, "2.2.8 Abdella et al. (2025): Majority Voting Ensemble "
            "of Seven CNNs")
    para(doc,
         "Abdella et al. evaluated seven pretrained CNN architectures "
         "(including GoogLeNet and Inception-v3) for four-class brain "
         "tumor classification on T1-weighted MRI images. Each model "
         "was trained using two optimisers (SGDM and Adam) and "
         "evaluated on a public dataset split into training (70%), "
         "validation (10%), and testing (20%) subsets. A majority "
         "voting ensemble aggregating predictions from all 14 trained "
         "models achieved 99.8% accuracy with AUC values above 0.997 "
         "for all tumor classes, outperforming every standalone model. "
         "The work highlights that model diversity is the key "
         "ingredient for successful ensembling, a finding that aligns "
         "with our project's deliberate architectural diversity "
         "(CNN + EfficientNet + Swin) rather than combining variants "
         "of the same family [9].",
         align="justify")

    heading(doc, 3, "2.2.9 Ali et al. (2024): Grad-CAM with ResNet50 for "
            "Clinical Explainability")
    para(doc,
         "Ali et al. addressed the critical challenge of "
         "interpretability in deep-learning brain tumor detection by "
         "integrating Grad-CAM with ResNet50. The methodology trained "
         "ResNet50 on an MRI dataset with data augmentation, then "
         "applied Grad-CAM to generate heatmap visualisations "
         "highlighting regions of the MRI most influential to the "
         "model's predictions. The model achieved 98.52% testing "
         "accuracy with precision-recall metrics exceeding 98%. The "
         "work demonstrates that combining high accuracy with visual "
         "explanations significantly increases clinical trust and aids "
         "radiologists in verifying that the model focuses on relevant "
         "pathological markers. Our project adopts the same Grad-CAM "
         "approach, extended to all three base models and the "
         "predicted class of the ensemble [10].",
         align="justify")

    heading(doc, 3, "2.2.10 Naser and Deen (2023): Multi-CAM Comparative "
            "Study")
    para(doc,
         "Naser and Deen developed a lightweight Explainable Deep "
         "Learning framework comparing three explainability "
         "techniques, Class Activation Mapping (CAM), Grad-CAM, and "
         "Grad-CAM++, applied to three models: pretrained VGG-19, "
         "scratch VGG-19, and EfficientNet. Evaluation on two "
         "benchmark MRI datasets (multi-class and binary) showed that "
         "pretrained VGG-19 with Grad-CAM outperformed other "
         "combinations both in classification accuracy and in the "
         "visual quality of localisation heatmaps. This study is one "
         "of the few that quantitatively compares different CAM "
         "variants for medical imaging, and confirms that Grad-CAM "
         "(which we use in our project) consistently produces the "
         "most clinically meaningful heatmaps [11].",
         align="justify")

    heading(doc, 2, "2.3 Summary of Reviewed Approaches")
    para(doc,
         "Table 2.1 consolidates the ten reviewed papers, showing for "
         "each one the method, dataset, headline accuracy, and the "
         "primary way in which the paper informed our project's "
         "design.",
         align="justify")
    para(doc, "Table 2.1: Summary comparison of the reviewed literature.",
         bold=True)
    add_table(doc,
              ["Reference", "Year", "Method",
               "Accuracy", "Contribution to our project"],
              [["Babu Vimala et al.", "2023", "EfficientNetB0-B4 + Grad-CAM",
                "99.06%", "Two-step transfer learning baseline"],
               ["Islam et al.", "2024", "EfficientNetB0-B7 with augmentation",
                "99.69%", "Progressive unfreezing schedule"],
               ["Alnowami et al.", "2024", "EfficientNetv2 + GAM + ECA",
                "99.76%", "Attention can rival transformers"],
               ["Liu et al.", "2021",
                "Swin Transformer (foundational)",
                "81.3% (ImageNet)", "Shifted-window self-attention"],
               ["Haq et al.", "2024",
                "Modified Swin + HSW-MSA + ResMLP",
                "99.92%", "Transformer-on-MRI feasibility"],
               ["Alsubai et al.", "2024",
                "Swin + ResNet50V2 hybrid",
                "high", "CNN + transformer fusion idea"],
               ["Khan et al.", "2023",
                "Weighted ensemble VGG19 + CNN",
                "best of 3", "Grid-search ensemble weights"],
               ["Abdella et al.", "2025",
                "Majority voting (7 CNNs)",
                "99.8%", "Diversity drives ensembles"],
               ["Ali et al.", "2024", "ResNet50 + Grad-CAM",
                "98.52%", "Clinical-trust framework"],
               ["Naser & Deen", "2023",
                "VGG-19 + CAM / Grad-CAM / Grad-CAM++",
                "high", "Grad-CAM beats CAM and Grad-CAM++"]])

    heading(doc, 2, "2.4 Market Research")
    para(doc,
         "The global burden of brain tumors is substantial and growing. "
         "According to the Global Cancer Observatory (GLOBOCAN 2022), "
         "approximately 308,000 new cases of brain and central nervous "
         "system tumors were diagnosed worldwide in 2020, with an "
         "age-standardised global incidence rate of 5.57 per 100,000 "
         "population. Brain tumors are the 19th most common cancer "
         "globally and the 12th leading cause of cancer death [17].",
         align="justify")
    para(doc,
         "In India, the picture is particularly acute. The International "
         "Association of Cancer Registries and the Indian Journal of "
         "Neurology (2025) report 28,000 new brain tumor cases per year, "
         "though some expert estimates range as high as 40,000 to 50,000 "
         "cases annually. Brain tumors are the second most common cancer "
         "site in children aged 0 to 14 years, after lymphoid leukaemia, "
         "according to the National Cancer Registry Programme (NCRP). The "
         "CNS tumor incidence rate in India ranges from 5 to 10 per 100,000 "
         "population [17].",
         align="justify")
    para(doc,
         "Against this disease burden, India faces a severe shortage of "
         "diagnostic specialists. The country has fewer than one radiologist "
         "per 100,000 people, with the deficit being most acute in tier-2 "
         "and tier-3 cities and in rural regions where expensive imaging "
         "machines often sit idle because there is no qualified personnel "
         "to read the scans. In this context, AI-assisted diagnostic tools "
         "that can provide reliable second-reader opinions in real time are "
         "not a luxury but a clinical necessity [17].",
         align="justify")
    para(doc,
         "The global medical image analysis market, driven by advances in "
         "deep learning and the digitisation of radiology workflows, is "
         "projected to grow significantly through 2030. Governments across "
         "the Asia-Pacific region, including India (through initiatives like "
         "Ayushman Bharat) and China (Healthy China 2030), are prioritising "
         "access to diagnostic imaging and AI-assisted interpretation tools "
         "as part of their public health strategies.",
         align="justify")

    heading(doc, 2, "2.5 Research Gap and Limitations of Existing Systems")
    para(doc,
         "Despite the substantial body of work, several limitations persist "
         "in the existing literature:",
         align="justify")
    bullet(doc,
           "Overstated accuracy: Many studies report only validation-set "
           "accuracy without a truly held-out test set, leading to inflated "
           "headline numbers. The val-to-test accuracy drop, which we "
           "observe to be 4 to 8 percentage points across our models, is "
           "rarely discussed.")
    bullet(doc,
           "Limited architectural diversity in ensembles: Most ensemble "
           "studies combine variants of the same architecture family "
           "(for example, multiple ResNets or multiple EfficientNets). "
           "Few works deliberately combine architecturally distinct model "
           "families (CNN, depthwise-separable CNN, and transformer) to "
           "maximise error diversity.")
    bullet(doc,
           "Insufficient explainability: While Grad-CAM is mentioned in "
           "several papers, few provide per-model heatmap comparisons or "
           "analyse failure modes through the lens of explainability.")
    bullet(doc,
           "Lack of deployment engineering: Most published works stop at "
           "model training and metric reporting. Few provide a working "
           "web interface, REST API, or containerised deployment that a "
           "hospital IT team could realistically adopt.")
    bullet(doc,
           "Absence of honest negative results: The ensemble learning "
           "literature consistently reports ensembles outperforming "
           "individual models. The scenario where a dominant base learner "
           "makes ensembling counterproductive is rarely acknowledged.")

    heading(doc, 2, "2.6 Motivation for Project Selection")
    para(doc,
         "The selection of this project was motivated by the convergence of "
         "a pressing clinical need (reliable brain tumor classification "
         "support), a well-characterised public dataset, and a desire to "
         "address the research gaps identified above. Specifically, the "
         "project seeks to answer the following questions:",
         align="justify")
    bullet(doc,
           "Can an ensemble of three architecturally diverse models (a "
           "custom CNN, a depthwise-separable CNN via transfer learning, "
           "and a vision transformer) yield better classification than any "
           "single model alone?")
    bullet(doc,
           "How large is the gap between validation-set and test-set "
           "accuracy, and what does this gap imply about the reliability "
           "of commonly reported accuracy figures in the literature?")
    bullet(doc,
           "Can Grad-CAM heatmaps provide clinically meaningful visual "
           "explanations, and do different model architectures attend to "
           "different anatomical regions?")
    bullet(doc,
           "Can the entire system, from image upload to prediction with "
           "explainability, be packaged as a deployable web application "
           "suitable for clinical demonstration?")
    para(doc,
         "By addressing these questions, the project contributes not only "
         "a technical artifact (the trained models and deployment pipeline) "
         "but also empirical insights, particularly the documentation of "
         "an honest negative ensembling result, that have pedagogical and "
         "practical value.",
         align="justify")

    doc.add_page_break()

    # ============================================================ CHAPTER 3
    heading(doc, 1, "Chapter 3: Dataset and Data Preparation")

    heading(doc, 2, "3.1 Overview")
    para(doc,
         "A well-characterised and reproducible dataset is the foundation "
         "of any supervised deep-learning system, and the credibility of "
         "every accuracy figure reported in this report rests directly on "
         "the quality, balance, and provenance of the data used to train "
         "and evaluate the models. This chapter describes the dataset "
         "source, the clinical meaning of each class, the rationale for "
         "selecting this dataset, the exploratory analysis performed, the "
         "train/validation/test split, the preprocessing pipeline, and the "
         "augmentation strategy.",
         align="justify")

    heading(doc, 2, "3.2 Dataset Source and Acquisition")
    para(doc,
         "The dataset used throughout this project is the publicly "
         "available Brain Tumor MRI Dataset, released by Masoud Nickparvar "
         "on the Kaggle platform. The dataset is freely accessible at:",
         align="justify")
    para(doc,
         "https://www.kaggle.com/datasets/masoudnickparvar/"
         "brain-tumor-mri-dataset",
         bold=True)
    para(doc,
         "The collection is a curated aggregation of three earlier public "
         "brain MRI resources: the figshare brain tumor dataset, the "
         "SARTAJ brain tumor dataset, and the Br35H brain tumor detection "
         "dataset, re-organised by the author into a single unified "
         "four-class taxonomy with a consistent Training/Testing folder "
         "layout. This aggregation of multiple upstream sources yields a "
         "heterogeneous corpus spanning multiple scanners, intensity "
         "profiles, and acquisition protocols [12].",
         align="justify")
    para(doc,
         "The dataset is distributed as plain image files (JPEG) organised "
         "in a two-level split/class/ directory layout. No DICOM headers, "
         "patient identifiers, scanner metadata, or other forms of "
         "protected health information accompany the images; each file is "
         "a single 2D axial brain MRI slice anonymised at source. After "
         "per-class balancing by the original author, the dataset contains "
         "a total of 7,200 images split as 5,600 in the Training folder "
         "and 1,600 in the Testing folder.",
         align="justify")

    heading(doc, 2, "3.3 Class Categories and Clinical Meaning")
    para(doc,
         "The four classes in the dataset correspond to the three most "
         "clinically common categories of primary brain tumor and a "
         "no-tumor control class.",
         align="justify")
    para(doc,
         "Glioma: Gliomas are neoplasms that arise from the glial cells "
         "of the central nervous system, predominantly astrocytes, "
         "oligodendrocytes, and ependymal cells. They constitute the most "
         "common form of malignant primary brain tumor in adults. The WHO "
         "grades gliomas on a scale from I (least aggressive) to IV (most "
         "aggressive), with glioblastoma multiforme (grade IV) being among "
         "the most aggressive of all human cancers. On T1-weighted MRI, "
         "gliomas frequently appear as ill-defined intra-axial masses, "
         "often with surrounding vasogenic oedema and internal "
         "heterogeneity.",
         align="justify")
    para(doc,
         "Meningioma: Meningiomas arise from the arachnoid cap cells of "
         "the meninges, the three layers of protective membrane surrounding "
         "the brain and spinal cord. They are the most common primary brain "
         "tumor overall and are typically benign and slow-growing. On MRI, "
         "meningiomas characteristically appear as well-circumscribed, "
         "extra-axial masses attached to the dura, often exhibiting the "
         "dural tail enhancement pattern. Despite their distinctive "
         "signature on multi-modal imaging, their appearance on a single "
         "T1-weighted axial slice can be visually similar to that of a "
         "glioma, making them the primary source of inter-class confusion.",
         align="justify")
    para(doc,
         "Pituitary tumor: Pituitary adenomas are growths arising from the "
         "pituitary gland, a small endocrine organ situated at the base of "
         "the brain in the sella turcica. The overwhelming majority are "
         "benign. On MRI, pituitary lesions appear within or just superior "
         "to the sella turcica, at the very base of the brain. This "
         "anatomically constrained location makes them the easiest of the "
         "three tumor types for an automated classifier to identify.",
         align="justify")
    para(doc,
         "No tumor: This class comprises MRI slices with no detectable "
         "mass lesion. Its inclusion is essential for clinical usability: "
         "a diagnostic system must be capable of confidently identifying a "
         "normal scan, because false-positive tumor predictions carry their "
         "own clinical costs, including patient anxiety and unnecessary "
         "follow-up procedures.",
         align="justify")

    heading(doc, 2, "3.4 Rationale for Dataset Selection")
    para(doc,
         "Several brain MRI datasets are available in the public domain, "
         "including the BraTS challenge collections, the standalone "
         "figshare dataset, the IXI dataset, and various "
         "single-institution releases. The Nickparvar dataset was selected "
         "after evaluating each candidate against the following criteria:",
         align="justify")
    bullet(doc,
           "Four-class taxonomy aligned with the clinical problem. Many "
           "alternative datasets address only binary tumor/no-tumor "
           "classification, or focus on volumetric segmentation rather "
           "than slice-level classification.")
    bullet(doc,
           "Class balance. After per-class augmentation by the original "
           "author, the Training folder contains a near-uniform number of "
           "images per class, which simplifies training and makes accuracy "
           "a more informative headline metric.")
    bullet(doc,
           "Sufficient volume for transfer learning. With approximately "
           "7,200 images, the dataset is large enough to meaningfully "
           "fine-tune ImageNet-pretrained backbones without entering the "
           "very-low-data regime.")
    bullet(doc,
           "Public and permissive licensing. The dataset contains no "
           "patient identifiers and is released for academic use.")
    bullet(doc,
           "Established benchmark. The dataset has been used in dozens of "
           "peer-reviewed studies since 2021, providing a competitive "
           "baseline against which our results can be compared.")
    bullet(doc,
           "Heterogeneity. Because the dataset aggregates three upstream "
           "collections, the images span multiple scanners and intensity "
           "profiles, forcing any model to be robust to inter-scanner "
           "variation.")

    heading(doc, 2, "3.5 Dataset Statistics")
    para(doc,
         "The full dataset contains the per-class counts shown in "
         "Table 3.1. The Training folder is partitioned via stratified "
         "random sampling into model-training and validation subsets at an "
         "80/20 ratio; the Testing folder is retained as a strictly "
         "held-out evaluation set.",
         align="justify")
    para(doc, "Table 3.1: Per-class image counts.", bold=True)
    add_table(doc,
              ["Split", "glioma", "meningioma", "notumor", "pituitary", "Total"],
              [["Training (full)", "1,400", "1,400", "1,400", "1,400", "5,600"],
               ["  Train (80%)", "1,120", "1,120", "1,120", "1,120", "4,480"],
               ["  Validation (20%)", "280", "280", "280", "280", "1,120"],
               ["Test (held out)", "400", "400", "400", "400", "1,600"],
               ["Grand total", "1,800", "1,800", "1,800", "1,800", "7,200"]])

    heading(doc, 2, "3.6 Exploratory Data Analysis")
    para(doc,
         "Before any model was trained, the dataset was characterised "
         "across four dimensions to identify anomalies requiring special "
         "handling in the preprocessing pipeline.",
         align="justify")
    para(doc,
         "File integrity: Every image file in both splits was opened with "
         "the Python Imaging Library. No corrupt files, truncated streams, "
         "or unreadable headers were encountered.",
         align="justify")
    para(doc,
         "Image modes: The dataset is heterogeneous in colour mode. Of "
         "the 7,200 images, 2,585 are stored as RGB, 1,892 as "
         "single-channel grayscale, and 3 as four-channel RGBA. A uniform "
         "conversion to RGB is applied at the start of every preprocessing "
         "pipeline.",
         align="justify")
    para(doc,
         "Pixel dimensions: Image dimensions vary widely, from a minimum "
         "of 150 x 168 pixels to a maximum of 1,375 x 1,446 pixels, with "
         "a median of approximately 512 x 512. All images are resized to "
         "224 x 224 during preprocessing.",
         align="justify")
    para(doc,
         "Pixel-intensity distribution: Per-class pixel-intensity "
         "histograms overlap substantially across all four classes, "
         "confirming that no class is separable from the others by simple "
         "thresholding of pixel values. Deep feature learning is therefore "
         "genuinely required.",
         align="justify")

    heading(doc, 2, "3.7 Train / Validation / Test Split Strategy")
    para(doc,
         "A central methodological commitment of this project is the "
         "separation of the available data into three subsets whose roles "
         "are strictly enforced throughout the project lifecycle:",
         align="justify")
    bullet(doc,
           "Training set (4,480 images, 62.2% of the corpus): Used for "
           "gradient updates during model fitting.")
    bullet(doc,
           "Validation set (1,120 images, 15.6%): Used for early-stopping "
           "signals, hyperparameter selection, ensemble-weight grid search, "
           "and any other model-selection decisions.")
    bullet(doc,
           "Test set (1,600 images, 22.2%): Held out entirely until all "
           "hyperparameters and modelling choices have been frozen, then "
           "evaluated exactly once to produce the final results.")
    para(doc,
         "The train/validation split is performed by "
         "sklearn.model_selection.train_test_split with stratify=labels "
         "and random_state=42, guaranteeing both class balance and "
         "bit-exact reproducibility. The resulting partition is cached as "
         "CSV files at data/processed/ and version-controlled together "
         "with the source code.",
         align="justify")
    para(doc,
         "The validation-to-test accuracy gap observed across the "
         "experiments, between 4 and 8 percentage points across all three "
         "models, confirms that any system reporting only validation-set "
         "metrics would substantially overstate its real-world performance.",
         align="justify")

    heading(doc, 2, "3.8 Image Preprocessing Pipeline")
    para(doc,
         "A single deterministic preprocessing transform is applied to "
         "validation and test images, and to every image at inference time. "
         "The training pipeline applies the same transform plus an "
         "additional stochastic augmentation stage described in Section 3.9.",
         align="justify")
    para(doc, "Table 3.3: Image preprocessing transforms.", bold=True)
    add_table(doc,
              ["Step", "Transform", "Parameters"],
              [["1", "Convert to RGB", "PIL.Image.convert('RGB')"],
               ["2", "Resize", "224 x 224 px, bilinear interpolation"],
               ["3", "Convert to tensor", "float32, range [0, 1], channels-first"],
               ["4", "Normalise (ImageNet)", "mean=[0.485, 0.456, 0.406], "
                "std=[0.229, 0.224, 0.225]"]])
    para(doc,
         "The choice of 224 x 224 is dictated by the published ImageNet "
         "checkpoints used for transfer learning. Both EfficientNet-B0 and "
         "Swin-Tiny were pretrained at this resolution, and using a "
         "different target would invalidate the pretrained positional "
         "embeddings (in the case of Swin) or the spatial expectations of "
         "early convolutional layers. The use of ImageNet mean and standard "
         "deviation at step 4 ensures inputs lie in the distribution the "
         "pretrained weights expect.",
         align="justify")

    heading(doc, 2, "3.9 Data Augmentation Strategy")
    para(doc,
         "Augmentation is applied only to training images and only at "
         "training time. Validation and test images are never augmented. "
         "The augmentation policy is deliberately conservative; "
         "augmentations that would produce clinically implausible images "
         "are explicitly excluded.",
         align="justify")
    para(doc, "Table 3.4: Training-time augmentation policy.", bold=True)
    add_table(doc,
              ["Augmentation", "Probability / Range", "Clinical Justification"],
              [["Random horizontal flip", "p = 0.5",
                "Brain is bilaterally symmetric on axial slices"],
               ["Random rotation", "+/- 15 degrees",
                "Simulates patient head tilt"],
               ["Random affine (translate)", "+/- 5%",
                "Simulates field-of-view variation"],
               ["Random affine (scale)", "+/- 5%",
                "Simulates magnification differences"],
               ["Colour jitter (brightness)", "+/- 0.20",
                "Simulates scanner-gain variation"],
               ["Colour jitter (contrast)", "+/- 0.20",
                "Simulates intensity-window variation"],
               ["Random erasing", "p=0.25, area 2-10%",
                "Forces model to reason about global structure"]])
    para(doc,
         "Three augmentations were considered and deliberately excluded. "
         "Vertical flips would invert cranio-caudal anatomy. Aggressive "
         "rotations beyond 30 degrees would produce unrealistic images. "
         "Hue jitter was found in pilot experiments to degrade validation "
         "accuracy because pixel intensities in MRI already carry "
         "diagnostic meaning.",
         align="justify")

    heading(doc, 2, "3.10 Ethical Considerations and Licensing")
    para(doc,
         "The dataset is publicly distributed for academic and research "
         "use and contains no patient-identifiable information. No "
         "additional patient data was collected during this project. The "
         "output of the trained system is intended for academic "
         "demonstration only and is not represented as a medical device or "
         "a substitute for the judgement of a qualified radiologist. The "
         "Streamlit interface displays a persistent advisory to this "
         "effect, and the low-confidence threshold gate defaults to "
         "recommending radiologist review for any prediction below 85% "
         "confidence.",
         align="justify")

    doc.add_page_break()

    # ============================================================ CHAPTER 4
    heading(doc, 1, "Chapter 4: Project Planning")

    heading(doc, 2, "4.1 Overview")
    para(doc,
         "Project planning defines the structured roadmap for the "
         "development and completion of the Brain Tumor Detection system. "
         "The planning process focuses on dividing the project into "
         "well-defined phases to ensure systematic progress, effective "
         "time management, and clear deliverables at each stage.",
         align="justify")

    heading(doc, 2, "4.2 Current Status of the Project")
    para(doc,
         "At the present stage, the project is fully implemented and "
         "operational. The following activities have been successfully "
         "completed:",
         align="justify")
    bullet(doc, "Identification of the problem domain and formulation of "
           "project objectives")
    bullet(doc, "Detailed literature survey and analysis of existing systems")
    bullet(doc, "Dataset acquisition, exploratory data analysis, and "
           "preprocessing pipeline design")
    bullet(doc, "Design and training of three deep learning models: custom "
           "CNN, EfficientNet-B0, and Swin-Tiny")
    bullet(doc, "Ensemble construction with grid-search weight optimisation")
    bullet(doc, "Rigorous evaluation on the held-out test set (1,600 images)")
    bullet(doc, "Grad-CAM explainability implementation for all three models")
    bullet(doc, "Streamlit web interface and FastAPI REST backend development")
    bullet(doc, "ONNX model export with inference verification")
    bullet(doc, "Docker containerisation for cloud deployment")
    bullet(doc, "Complete documentation and report preparation")

    heading(doc, 2, "4.3 System Architecture")
    para(doc,
         "The Brain Tumor Detection system follows a layered architecture "
         "designed for modularity, maintainability, and ease of deployment. "
         "The system is divided into four logical layers:",
         align="justify")
    para(doc,
         "Presentation Layer: This layer provides two user-facing "
         "interfaces. The Streamlit web application offers an interactive "
         "demo experience with drag-and-drop image upload, real-time "
         "prediction display, confidence metrics, per-model probability "
         "breakdowns, and Grad-CAM heatmap overlays. The FastAPI REST "
         "backend exposes the same prediction capability as a JSON API, "
         "suitable for integration with hospital PACS systems, mobile "
         "applications, or third-party scripts.",
         align="justify")
    para(doc,
         "Application Layer: The core processing logic resides in the "
         "src/ Python package. The BrainTumorPredictor class in "
         "src/inference.py is the single entry point for all prediction "
         "requests. It loads all three trained models once, runs each "
         "model on the input image, combines their outputs via the "
         "ensemble combiner, and optionally generates Grad-CAM heatmaps. "
         "Both the Streamlit app and the FastAPI backend import and call "
         "this same class, ensuring consistent behaviour across interfaces.",
         align="justify")
    para(doc,
         "Model Layer: This layer contains the three trained model "
         "architectures (src/cnn_model.py, src/transfer_model.py, "
         "src/swin_model.py), the ensemble combiner (src/ensemble.py), "
         "the Grad-CAM implementation (src/explainability.py), and the "
         "training engine (src/train.py). All model weights are stored "
         "as .pth checkpoint files in the models/ directory, alongside "
         "an ensemble_config.json that records the ensemble method and "
         "optimal weights.",
         align="justify")
    para(doc,
         "Data Layer: The dataset itself resides in DATASET/ with the "
         "Training/ and Testing/ folders. Cached stratified split CSVs "
         "are stored in data/processed/. Training outputs (plots, "
         "confusion matrices, Grad-CAM images, CSV reports) are written "
         "to outputs/. TensorBoard event files are logged to logs/.",
         align="justify")

    heading(doc, 2, "4.4 Identified Challenges and System Limitations")
    bullet(doc,
           "GPU memory constraint: The entire training pipeline must fit "
           "within 4 GB of VRAM (NVIDIA GTX 1650). This limits the maximum "
           "batch size to 32 for Swin-Tiny under mixed precision.")
    bullet(doc,
           "Mixed-precision numerical instability: fp16 training triggered "
           "a NaN bug in cuDNN on the Turing GPU architecture, requiring a "
           "switch to bfloat16.")
    bullet(doc,
           "Dominant-model effect: One base model (EfficientNet-B0) is "
           "significantly stronger than the other two, limiting the "
           "effectiveness of ensembling.")
    bullet(doc,
           "Single MRI sequence: The system classifies from a single "
           "T1-weighted slice, whereas clinical decisions typically combine "
           "multiple MRI sequences.")

    heading(doc, 2, "4.5 Strategy to Address Identified Challenges")
    bullet(doc,
           "Bfloat16 mixed precision: Replacing fp16 with bfloat16 "
           "resolved the NaN issue while retaining the speed benefits of "
           "mixed-precision training.")
    bullet(doc,
           "Aggressive but clinically valid augmentation: Seven "
           "augmentation transforms effectively multiply the training "
           "data volume, mitigating overfitting within the limited dataset.")
    bullet(doc,
           "Architectural diversity: Using three fundamentally different "
           "architectures (vanilla CNN, depthwise-separable CNN, and "
           "transformer) maximises the potential for uncorrelated errors.")
    bullet(doc,
           "Grad-CAM integration: Providing visual explanations alongside "
           "every prediction addresses the clinical-trust barrier.")
    bullet(doc,
           "Threshold gate: A configurable confidence threshold flags "
           "uncertain predictions for radiologist review rather than "
           "presenting them with false confidence.")

    heading(doc, 2, "4.6 Implementation Roadmap")
    para(doc, "Table 4.1: Phases of the entire project work.", bold=True)
    add_table(doc,
              ["Phase", "Module", "Tasks"],
              [["1", "Data Acquisition and EDA",
                "Download dataset, verify integrity, compute class "
                "distributions, analyse image modes and dimensions, "
                "build stratified train/val/test split"],
               ["2", "Custom CNN Training",
                "Design 4-block VGG-style CNN, implement training "
                "engine with AMP and early stopping, train on 4,480 "
                "images, evaluate on validation set"],
               ["3", "Transfer Learning (EfficientNet-B0)",
                "Load ImageNet-pretrained backbone, replace classifier "
                "head, two-stage training (freeze then fine-tune), "
                "evaluate on validation set"],
               ["4", "Swin Transformer Fine-tuning",
                "Load timm Swin-Tiny checkpoint, replace head, "
                "two-stage training with very small Stage 2 LR, "
                "evaluate on validation set"],
               ["5", "Ensemble Construction",
                "Implement soft voting, weighted average, and stacking "
                "combiners, run exhaustive grid search for optimal "
                "weights, select best method"],
               ["6", "Test-Set Evaluation",
                "Run all three models and ensemble on held-out 1,600 "
                "test images, compute accuracy, F1, AUC, confusion "
                "matrices, ROC curves"],
               ["7", "Explainability (Grad-CAM)",
                "Implement Grad-CAM for all three architectures, "
                "generate heatmap overlays, analyse failure modes"],
               ["8", "Deployment",
                "Build Streamlit UI, FastAPI REST endpoint, ONNX "
                "export pipeline, Docker container"],
               ["9", "Documentation",
                "Write project report, prepare viva QA guide, "
                "generate flow diagrams and architecture documentation"]])

    heading(doc, 2, "4.7 Project Scheduling and Gantt Chart")
    para(doc,
         "The project was executed over a single academic semester. "
         "The Gantt chart (Figure 4.2) illustrates the timeline and "
         "sequencing of activities. Phases 2 through 4 (model training) "
         "were partially parallelised, with the custom CNN trained first "
         "to establish a baseline, followed by the two transfer-learning "
         "models. The ensemble and evaluation phases were strictly "
         "sequential, as they depend on the availability of all three "
         "trained models.",
         align="justify")

    doc.add_page_break()

    # ============================================================ CHAPTER 5
    heading(doc, 1, "Chapter 5: Project Description")

    heading(doc, 2, "5.1 Software Model")
    para(doc,
         "The Brain Tumor Detection system is built using a modular "
         "Python package architecture. The src/ directory contains all "
         "production code, organised into single-responsibility modules. "
         "The app/ directory contains the two user-facing interfaces. "
         "The notebooks/ directory contains seven Jupyter notebooks that "
         "reproduce the entire experimental pipeline with embedded outputs "
         "for auditability.",
         align="justify")
    para(doc,
         "The key technology choices and their justifications are "
         "summarised below:",
         align="justify")
    add_table(doc,
              ["Component", "Choice", "Reason"],
              [["Framework", "PyTorch 2.12", "Industry standard for research; "
                "native CUDA support"],
               ["Mixed precision", "bfloat16", "Avoids fp16 cuDNN NaN bug "
                "on Turing GPUs"],
               ["Optimiser", "AdamW", "Strong baseline for small medical "
                "datasets"],
               ["Loss function", "CrossEntropyLoss", "Classes are balanced; "
                "no weighting needed"],
               ["Batch size", "32", "Largest that fits Swin-Tiny in 4 GB "
                "VRAM with AMP"],
               ["LR scheduler", "CosineAnnealingLR", "Smooth convergence "
                "with minimal tuning"],
               ["Random seed", "42 (all RNGs)", "Bit-exact reproducibility"],
               ["Hardware", "NVIDIA GTX 1650 (4 GB), Windows 10, "
                "Python 3.11", ""]])

    heading(doc, 2, "5.2 Software Requirements Specification (SRS)")

    heading(doc, 3, "5.2.1 Introduction")
    para(doc,
         "The Brain Tumor Detection system is a deep-learning-based "
         "diagnostic support tool that classifies brain MRI slices into "
         "four categories. This SRS document outlines the functional and "
         "non-functional requirements, system constraints, and expected "
         "performance characteristics.",
         align="justify")

    heading(doc, 3, "5.2.2 General Description")
    para(doc,
         "The system is a client-server application with a Python backend "
         "and a web-based frontend. It supports two modes of interaction: "
         "(1) a Streamlit web interface for interactive single-image "
         "prediction with visual feedback, and (2) a FastAPI REST endpoint "
         "for programmatic batch prediction. Both interfaces wrap the same "
         "BrainTumorPredictor class, ensuring identical results regardless "
         "of the access method.",
         align="justify")

    heading(doc, 3, "5.2.3 Functional Requirements")
    para(doc, "The system shall:", align="justify")
    bullet(doc, "Accept a brain MRI image in JPG, PNG, or BMP format as "
           "input, with no constraint on the original image dimensions.")
    bullet(doc, "Preprocess the input image by converting to RGB, resizing "
           "to 224 x 224, and normalising with ImageNet statistics.")
    bullet(doc, "Run the preprocessed image through all three trained "
           "models (custom CNN, EfficientNet-B0, Swin-Tiny) and return "
           "softmax probability vectors from each.")
    bullet(doc, "Combine the three probability vectors using the saved "
           "ensemble method and weights to produce a single ensemble "
           "probability distribution.")
    bullet(doc, "Return the predicted class (the class with the highest "
           "ensemble probability) and the corresponding confidence score.")
    bullet(doc, "Optionally generate Grad-CAM heatmap overlays for each "
           "base model, targeted at the predicted class.")
    bullet(doc, "Flag predictions whose confidence falls below a "
           "user-configurable threshold (default 0.85) with a "
           '"consult radiologist" advisory.')
    bullet(doc, "Display results in the Streamlit UI including: the "
           "predicted class with colour coding, the confidence score, "
           "a probability bar chart, a per-model probability table, and "
           "three Grad-CAM heatmap images.")
    bullet(doc, "Expose a /predict REST endpoint that accepts a "
           "multipart/form-data image upload and returns a JSON response "
           "with the predicted class, confidence, and per-model "
           "probabilities.")

    heading(doc, 3, "5.2.4 Interface Requirements")
    para(doc, "User Interface:", bold=True)
    bullet(doc, "Web-based interface accessible via any modern browser.")
    bullet(doc, "Drag-and-drop image upload or dropdown-based test-sample "
           "selection.")
    bullet(doc, "Sidebar with model information and configurable settings.")
    para(doc, "Hardware Interface:", bold=True)
    bullet(doc, "Runs on standard desktop or laptop with or without a "
           "dedicated GPU. Falls back gracefully to CPU.")
    para(doc, "Software Interface:", bold=True)
    bullet(doc, "Python 3.11+, PyTorch 2.12, timm, Streamlit, FastAPI, "
           "ONNX Runtime.")

    heading(doc, 3, "5.2.5 Performance Requirements")
    bullet(doc, "Single-image inference shall complete within 300 ms on "
           "GPU (including Grad-CAM) or within 1 second on CPU.")
    bullet(doc, "Model loading shall complete within 10 seconds on first "
           "startup.")
    bullet(doc, "The Streamlit interface shall remain responsive during "
           "prediction.")

    heading(doc, 3, "5.2.6 Design Constraints")
    bullet(doc, "The system shall be developed using open-source "
           "frameworks and tools.")
    bullet(doc, "All model weights shall be stored locally; no external "
           "API calls are required at inference time.")
    bullet(doc, "The system must be deployable on standard institutional "
           "infrastructure or personal laptops.")

    heading(doc, 3, "5.2.7 Non-Functional Attributes")
    para(doc,
         "Reliability: The system shall ensure reliable and deterministic "
         "predictions for the same input image. Reproducibility is "
         "guaranteed by fixed random seeds and cached data splits.",
         align="justify")
    para(doc,
         "Security: No user data is stored or transmitted externally. "
         "Images are processed in memory and discarded after prediction.",
         align="justify")
    para(doc,
         "Maintainability: Modular src/ package structure enables easy "
         "updates. Each model can be retrained and replaced independently "
         "without affecting other components.",
         align="justify")
    para(doc,
         "Scalability: The FastAPI backend supports concurrent requests. "
         "ONNX exports enable deployment on edge devices without PyTorch.",
         align="justify")

    heading(doc, 2, "5.3 Functional Specification")

    heading(doc, 3, "5.3.1 Deep Learning Models")
    para(doc,
         "Model 1: Custom CNN (src/cnn_model.py)",
         bold=True, align="justify")
    para(doc,
         "A VGG-style architecture consisting of four convolutional blocks. "
         "Each block contains two stacked 3 x 3 convolutions, each "
         "followed by Batch Normalisation and ReLU activation, and a 2 x 2 "
         "max-pooling layer that halves the spatial dimensions. The channel "
         "count doubles at each block: 32, 64, 128, 256. After the fourth "
         "block, the feature map (256 x 14 x 14) passes through Global "
         "Average Pooling, which collapses the spatial dimensions to "
         "produce a 256-dimensional feature vector. This vector is fed "
         "through a two-layer fully-connected classifier: a Dropout layer "
         "(p=0.5) followed by a Linear layer (256 to 128) with ReLU, then "
         "another Dropout (p=0.3) and a final Linear layer (128 to 4) that "
         "outputs four logits. Weight initialisation uses Kaiming normal "
         "(mode=fan_out, ReLU). Total parameters: 1,206,628.",
         align="justify")

    para(doc,
         "Model 2: EfficientNet-B0 Transfer Learning (src/transfer_model.py)",
         bold=True, align="justify")
    para(doc,
         "EfficientNet-B0 was chosen over ResNet-50, EfficientNet-B3, "
         "DenseNet-121, and ConvNeXt-Tiny on the basis of its "
         "accuracy-per-parameter efficiency (77.7% ImageNet top-1 at only "
         "5.3 million parameters) and its compatibility with the 4 GB VRAM "
         "budget. The pretrained backbone is loaded from torchvision; its "
         "1000-class classifier head is replaced with a new head: "
         "Dropout(0.3) followed by Linear(1280 to 4).",
         align="justify")
    para(doc,
         "Training proceeds in two stages. Stage 1 (warm-up, 4 epochs, "
         "learning rate = 0.001): The backbone is frozen and only the "
         "5,124 head parameters receive gradients. The module's .train() "
         "method is overridden to keep the frozen backbone in .eval() mode, "
         "preventing BatchNorm running-statistic drift. Stage 2 (fine-tune, "
         "8 epochs, learning rate = 0.0001): The backbone is unfrozen. "
         "The 10x smaller learning rate prevents catastrophic forgetting "
         "of pretrained features. Total parameters: 4,012,672.",
         align="justify")

    para(doc,
         "Model 3: Swin-Tiny Transformer (src/swin_model.py)",
         bold=True, align="justify")
    para(doc,
         "Swin-Tiny (swin_tiny_patch4_window7_224 from the timm library) "
         "is a hierarchical vision transformer that splits the input image "
         "into 4 x 4 patches and processes them through four stages of "
         "windowed self-attention blocks. Each stage reduces the spatial "
         "resolution via patch merging (analogous to pooling in CNNs) "
         "while increasing the feature dimensionality. The key innovation "
         "of Swin is the shifted-window mechanism: in alternating layers, "
         "the attention windows are shifted by half their size, allowing "
         "information to flow across window boundaries without the "
         "quadratic cost of global attention.",
         align="justify")
    para(doc,
         "The classifier head is replaced identically to EfficientNet "
         "(Dropout 0.3, Linear 768 to 4). Two-stage training mirrors the "
         "transfer-learning recipe but with a far smaller Stage 2 learning "
         "rate of 0.00002 (50x smaller than Stage 1) because transformer "
         "attention weights are extremely sensitive to large gradient "
         "updates. Total parameters: 27,522,430.",
         align="justify")

    para(doc, "Table 5.3: Architectural summary.", bold=True)
    add_table(doc,
              ["Property", "Custom CNN", "EfficientNet-B0", "Swin-Tiny"],
              [["Total params", "1.21 M", "4.01 M", "27.52 M"],
               ["Architecture type", "VGG-style conv blocks",
                "Depthwise-separable + SE", "Shifted-window self-attention"],
               ["Pretrained", "No (from scratch)", "Yes (ImageNet)",
                "Yes (ImageNet)"],
               ["Stage 1 LR", "0.001", "0.001", "0.001"],
               ["Stage 2 LR", "N/A (single stage)", "0.0001", "0.00002"],
               ["Stage 1 epochs", "20", "4", "3"],
               ["Stage 2 epochs", "N/A", "8", "4"]])

    heading(doc, 3, "5.3.2 Mathematical Model and Loss Functions")
    para(doc,
         "All three models are trained with the standard cross-entropy "
         "loss function. For a single sample with true class y and "
         "predicted probability p_y for that class, the loss is:",
         align="justify")
    para(doc, "    L = -log(p_y)", bold=True, align="center")
    para(doc,
         "where p_y is the softmax probability assigned to the correct "
         "class. Since the four classes are perfectly balanced, no "
         "class-weighting is applied to the loss.",
         align="justify")
    para(doc,
         "The optimiser used across all three models is AdamW "
         "(Adam with decoupled weight decay). The key hyperparameters "
         "are: beta_1 = 0.9, beta_2 = 0.999, weight decay = 0.0001. "
         "The learning rate schedule follows CosineAnnealingLR, which "
         "smoothly decays the learning rate from its initial value to "
         "eta_min = initial_lr x 0.01 over the course of training. "
         "Mixed-precision training uses bfloat16 via "
         "torch.amp.autocast, which halves memory usage and increases "
         "throughput on compatible GPUs without the numerical instability "
         "issues associated with fp16 on Turing-architecture GPUs.",
         align="justify")

    heading(doc, 3, "5.3.3 Ensemble Decision Logic")
    para(doc,
         "Each trained model outputs a softmax probability vector "
         "p_i in R^4 (one probability per class). Three ensemble "
         "combination methods were evaluated:",
         align="justify")
    para(doc,
         "Soft voting: The ensemble probability is the simple arithmetic "
         "mean of the three model outputs: p_ens = (p_cnn + p_tl + "
         "p_swin) / 3.",
         align="justify")
    para(doc,
         "Weighted average: The ensemble probability is a weighted sum: "
         "p_ens = w_cnn * p_cnn + w_tl * p_tl + w_swin * p_swin, where "
         "the weights w_i satisfy w_i >= 0 and w_cnn + w_tl + w_swin = 1. "
         "The optimal weights were found by exhaustive grid search over "
         "the simplex at step size 0.05, yielding 231 candidate weight "
         "combinations for three models. Each combination was evaluated "
         "on the validation set, and the best-performing combination was "
         "selected: [w_cnn=0.0, w_tl=0.9, w_swin=0.1].",
         align="justify")
    para(doc,
         "Stacking: A logistic regression meta-learner was trained on the "
         "concatenated 12-dimensional probability vector (4 probabilities "
         "from each of 3 models) using 5-fold cross-validation on the "
         "validation set.",
         align="justify")
    para(doc,
         "The grid search effectively discarded the custom CNN (w_cnn=0.0) "
         "because its predictions overlap too much with EfficientNet's. "
         "Adding the CNN brings no complementary signal but does add noise. "
         "The weighted average method was selected as the final ensemble "
         "combiner based on its validation-set performance.",
         align="justify")

    heading(doc, 3, "5.3.4 Grad-CAM Explainability Generation")
    para(doc,
         "Gradient-weighted Class Activation Mapping (Grad-CAM) is a "
         "technique that produces a coarse localisation heatmap "
         "highlighting the regions of an input image that were most "
         "important for a specific predicted class. It works by computing "
         "the gradient of the target class score with respect to the "
         "feature maps of a chosen convolutional layer, then using these "
         "gradients as weights for a linear combination of the feature "
         "maps [13].",
         align="justify")
    para(doc,
         "For the custom CNN and EfficientNet-B0, the target layer is the "
         "last convolutional layer before the global average pooling. For "
         "Swin-Tiny, which does not have traditional convolutional layers, "
         "the target is the LayerNorm at the end of the final transformer "
         "stage. A reshape operation is required because Swin's activations "
         "are in a channels-last format (sequence_length x channels) "
         "rather than the channels-first format (channels x height x width) "
         "that Grad-CAM expects. The implementation uses the "
         "pytorch-grad-cam library by Jacob Gildenblat.",
         align="justify")
    para(doc,
         "The resulting heatmap is a single-channel grayscale image where "
         "red indicates high importance and blue indicates low importance. "
         "This heatmap is overlaid on the original MRI image at 45% "
         "opacity to produce a composite visualisation that allows the "
         "clinician to see both the anatomical structure and the model's "
         "attention pattern simultaneously.",
         align="justify")

    heading(doc, 2, "5.4 Design Specification")

    heading(doc, 3, "5.4.1 Use-Case Diagrams")
    para(doc,
         "The system supports three primary actor roles:",
         align="justify")
    para(doc,
         "Radiologist / Clinician: Uploads a patient MRI image (or selects "
         "a test sample), views the predicted class and confidence, examines "
         "Grad-CAM heatmaps to verify the model's reasoning, and adjusts "
         "the confidence threshold.",
         align="justify")
    para(doc,
         "Researcher / Developer: Accesses the FastAPI endpoint for batch "
         "predictions, examines per-model probability breakdowns, runs "
         "evaluation scripts on custom datasets, and exports models to "
         "ONNX format.",
         align="justify")
    para(doc,
         "Administrator: Deploys the Docker container, configures server "
         "settings, monitors application logs, and manages model "
         "checkpoint updates.",
         align="justify")

    heading(doc, 3, "5.4.2 Data Flow Diagrams")
    para(doc,
         "Training Pipeline DFD: Raw MRI images flow from DATASET/ through "
         "the data loader (src/data_loader.py), which applies the "
         "preprocessing and augmentation transforms (src/preprocess.py). "
         "The processed batches flow into the training engine "
         "(src/train.py), which computes forward passes, calculates loss, "
         "performs backpropagation, and logs metrics to TensorBoard. At "
         "each epoch, the engine evaluates on the validation set and saves "
         "the best checkpoint to models/.",
         align="justify")
    para(doc,
         "Inference Pipeline DFD: A user uploads an image via the "
         "Streamlit UI or FastAPI endpoint. The image flows to "
         "BrainTumorPredictor (src/inference.py), which loads the image, "
         "applies the inference transform, runs all three models in "
         "parallel, combines their softmax outputs via the ensemble "
         "combiner, and optionally computes Grad-CAM overlays. The "
         "PredictionResult object is returned to the presentation layer "
         "for display.",
         align="justify")

    heading(doc, 3, "5.4.3 Data Dictionary")
    para(doc,
         "The following tables describe the key data structures used "
         "throughout the system.",
         align="justify")
    para(doc, "Table 5.6: Ensemble Configuration Schema "
         "(ensemble_config.json).", bold=True)
    add_table(doc,
              ["Field", "Type", "Description"],
              [["method", "string", "Ensemble method: 'weighted', "
                "'soft_voting', or 'stacking'"],
               ["weights", "list of float", "Mixing weights for the "
                "three models, summing to 1.0"],
               ["val_accuracy", "float", "Validation accuracy achieved "
                "with these weights"],
               ["models", "list of string", "Ordered list of model names: "
                "['cnn', 'transfer', 'swin']"]])

    para(doc, "Table 5.7: Prediction Result Schema.", bold=True)
    add_table(doc,
              ["Field", "Type", "Description"],
              [["predicted_class", "string", "One of: glioma, meningioma, "
                "pituitary, notumor"],
               ["predicted_class_idx", "int", "Integer index (0-3)"],
               ["confidence", "float", "Ensemble top-class probability"],
               ["ensemble_probs", "dict", "Class name to probability mapping"],
               ["model_probs", "dict of dict", "Per-model, per-class probs"],
               ["inference_time_ms", "float", "Wall-clock inference time"],
               ["gradcams", "dict or None", "Model name to heatmap array "
                "(if requested)"]])

    para(doc, "Table 5.8: Model Checkpoint Registry.", bold=True)
    add_table(doc,
              ["File", "Size", "Contents"],
              [["cnn_model.pth", "4.9 MB", "Custom CNN state_dict"],
               ["transfer_model.pth", "16.4 MB", "EfficientNet-B0 state_dict"],
               ["swin_model.pth", "110.2 MB", "Swin-Tiny state_dict"],
               ["cnn_model.onnx", "4.8 MB", "ONNX export of custom CNN"],
               ["transfer_model.onnx", "16.0 MB", "ONNX export of "
                "EfficientNet-B0"],
               ["swin_model.onnx", "112.9 MB", "ONNX export of Swin-Tiny"],
               ["ensemble_config.json", "0.4 KB", "Ensemble method and "
                "weights"]])

    para(doc, "Table 5.11: Inference API Schema (/predict).", bold=True)
    add_table(doc,
              ["Direction", "Field", "Type", "Description"],
              [["Request", "file", "multipart/form-data", "MRI image file"],
               ["Response", "predicted_class", "string", "Predicted category"],
               ["Response", "confidence", "float", "Top-class probability"],
               ["Response", "ensemble_probs", "object", "All 4 class probs"],
               ["Response", "model_probs", "object", "Per-model breakdown"],
               ["Response", "inference_time_ms", "float", "Processing time"]])

    doc.add_page_break()

    # ============================================================ CHAPTER 6
    heading(doc, 1, "Chapter 6: Implementation Issues")

    heading(doc, 2, "6.1 Dataset Variability and Image Quality")
    para(doc,
         "The most immediate challenge encountered during implementation "
         "was the heterogeneity of the source images. Because the "
         "Nickparvar dataset aggregates three separate upstream collections, "
         "the images vary in colour mode (RGB, grayscale, RGBA), "
         "resolution (150 x 168 to 1,375 x 1,446 pixels), intensity "
         "profile, and contrast characteristics. While this heterogeneity "
         "is ultimately beneficial for model robustness, it required "
         "careful preprocessing to ensure that every image entering the "
         "model pipeline was in a consistent format (3-channel, "
         "224 x 224, ImageNet-normalised). Edge cases such as the three "
         "RGBA images required explicit handling in the data loader.",
         align="justify")

    heading(doc, 2, "6.2 Model Selection and Hyperparameter Sensitivity")
    para(doc,
         "During the initial phase of model selection, five candidate "
         "transfer-learning backbones were evaluated: ResNet-50, "
         "EfficientNet-B0, EfficientNet-B3, DenseNet-121, and "
         "ConvNeXt-Tiny. EfficientNet-B0 was selected for its superior "
         "accuracy-per-parameter ratio and its ability to fit within the "
         "4 GB VRAM budget. EfficientNet-B3 achieved marginally higher "
         "validation accuracy but exceeded the VRAM limit at batch size "
         "32, requiring a batch size reduction to 16 that increased "
         "training time without meaningful accuracy improvement.",
         align="justify")
    para(doc,
         "The Swin Transformer was found to be extremely sensitive to the "
         "Stage 2 (fine-tuning) learning rate. A Stage 2 LR of 0.0001 "
         "(the value used for EfficientNet) caused the transformer's "
         "validation accuracy to collapse within three epochs. Reducing "
         "the Stage 2 LR to 0.00002 (50x smaller than Stage 1) resolved "
         "the issue. This sensitivity arises because the self-attention "
         "weights in a transformer encode global relationship patterns "
         "that are easily disrupted by large gradient updates.",
         align="justify")

    heading(doc, 2, "6.3 GPU Memory Constraints and Mixed Precision")
    para(doc,
         "The entire training pipeline was constrained to 4 GB of VRAM "
         "on an NVIDIA GTX 1650 (Turing architecture, compute capability "
         "7.5). Mixed-precision training using fp16 was initially employed "
         "to reduce memory usage and increase throughput. However, during "
         "Swin-Tiny training, the model produced NaN values in the loss "
         "after approximately 5 epochs.",
         align="justify")
    para(doc,
         "Root-cause analysis using layer-by-layer activation tracing "
         "identified the issue as a cuDNN numerical instability in the "
         "LayerNorm operations of the Swin Transformer when computed in "
         "fp16 on Turing GPUs. The solution was to switch to bfloat16 "
         "mixed precision, which maintains the same dynamic range as "
         "float32 (8 exponent bits) while using only 16 bits total. "
         "This switch resolved the NaN issue completely while retaining "
         "approximately 90% of the speed benefit of mixed-precision "
         "training.",
         align="justify")

    heading(doc, 2, "6.4 Class Confusion and Generalisation Gap")
    para(doc,
         "Analysis of the confusion matrices reveals that the primary "
         "source of classification error across all three models is the "
         "glioma-meningioma confusion pair. The custom CNN misclassifies "
         "17% of gliomas as meningiomas; EfficientNet-B0 reduces this to "
         "1%; Swin-Tiny to 6%. This confusion is clinically expected: on "
         "a single T1-weighted axial slice, these two tumor types can "
         "appear visually similar, and even experienced radiologists "
         "sometimes require additional MRI sequences to distinguish them.",
         align="justify")
    para(doc,
         "The generalisation gap between validation and test accuracy "
         "ranges from 4 to 8 percentage points across the three models. "
         "This gap is consistent with the known characteristics of small "
         "medical imaging datasets and underscores the importance of "
         "reporting test-set (not just validation-set) performance. Any "
         "system that reported only validation accuracy would overstate "
         "its real-world capability by a clinically significant margin.",
         align="justify")

    heading(doc, 2, "6.5 Clinical Adoption and Explainability Trust")
    para(doc,
         "From a practical deployment perspective, the most significant "
         "barrier to clinical adoption of AI diagnostic tools is not "
         "accuracy but trust. Clinicians need to understand why a model "
         "makes a particular prediction before they will act on it. The "
         "system addresses this through three mechanisms: (1) Grad-CAM "
         "heatmaps that show which regions of the MRI the model attended "
         "to, (2) per-model probability breakdowns that show whether the "
         "three models agree or disagree, and (3) a configurable "
         "confidence threshold that explicitly flags low-confidence "
         "predictions for radiologist review rather than presenting them "
         "with false confidence.",
         align="justify")
    para(doc,
         "Qualitative analysis of Grad-CAM outputs on correctly classified "
         "samples confirms that all three models attend to anatomically "
         "meaningful regions (the tumor mass for tumor classes, brain "
         "parenchyma for no-tumor). EfficientNet's heatmaps are the most "
         "spatially focused; Swin's are the most diffuse, consistent with "
         "the wider effective receptive field of self-attention. On "
         "misclassified samples, Grad-CAM reveals that the model often "
         "attends to the correct anatomical region but misinterprets the "
         "visual pattern, pointing to genuine clinical ambiguity rather "
         "than spurious cue exploitation.",
         align="justify")

    heading(doc, 2, "6.6 The Dominant-Model Problem in Ensembling")
    para(doc,
         "The most important and unexpected empirical finding of this "
         "project is that the three-model ensemble (94.69% test accuracy) "
         "does not outperform the best individual model (EfficientNet-B0, "
         "94.94% test accuracy). This result contradicts the common "
         "expectation that ensembles always improve over individual models, "
         "and it deserves careful analysis.",
         align="justify")
    para(doc,
         "The result is explained by two factors. First, EfficientNet-B0 "
         "already operates near the data's noise ceiling, with only 1.07% "
         "validation error and 5.06% test error. Improving on it would "
         "require base models that produce genuinely complementary errors. "
         "Second, the custom CNN and Swin-Tiny make many of the same "
         "errors as EfficientNet (particularly on the glioma-meningioma "
         "boundary), so averaging their predictions does not cancel error "
         "but instead injects noise from the weaker models.",
         align="justify")
    para(doc,
         "This phenomenon is documented in the ensemble learning "
         "literature as the 'dominant-model problem': when one base "
         "learner substantially outperforms the others, the variance "
         "reduction promised by ensembling only materialises if the base "
         "learners have uncorrelated errors. When the errors are "
         "correlated (as ours are, because all three models struggle on "
         "the same glioma-meningioma boundary), the ensemble provides no "
         "benefit and may even degrade performance.",
         align="justify")
    para(doc,
         "We report this finding honestly because it has both practical "
         "and pedagogical value. Practically, it demonstrates that "
         "ensemble construction requires careful consideration of error "
         "diversity, not just architectural diversity. Pedagogically, it "
         "provides a concrete example of a well-documented but rarely "
         "reported phenomenon in the academic literature, where positive "
         "results are preferentially published.",
         align="justify")

    para(doc, "Table 6.1: Per-model test-set performance.", bold=True)
    add_table(doc,
              ["Model", "Params", "Val acc", "Test acc", "Macro F1",
               "Macro AUC", "Train time", "Latency"],
              [["Custom CNN", "1.21 M", "0.9036", "0.8275", "0.8221",
                "0.9513", "32.1 min", "6.8 ms"],
               ["EfficientNet-B0", "4.01 M", "0.9893", "0.9494", "0.9482",
                "0.9908", "14.9 min", "24.1 ms"],
               ["Swin-Tiny", "27.52 M", "0.9670", "0.9194", "0.9171",
                "0.9865", "12.5 min", "25.7 ms"],
               ["Ensemble (weighted)", "32.74 M", "0.9893", "0.9469",
                "0.9456", "0.9907", "59.5 min", "56.6 ms"]])

    para(doc, "Table 6.2: Ensemble grid-search results (top 5).", bold=True)
    add_table(doc,
              ["Rank", "w_cnn", "w_transfer", "w_swin", "Val accuracy"],
              [["1", "0.00", "0.90", "0.10", "0.9893"],
               ["2", "0.00", "0.95", "0.05", "0.9893"],
               ["3", "0.00", "1.00", "0.00", "0.9893"],
               ["4", "0.05", "0.90", "0.05", "0.9884"],
               ["5", "0.00", "0.85", "0.15", "0.9875"]])

    doc.add_page_break()

    # ============================================================ CHAPTER 7
    heading(doc, 1, "Chapter 7: Conclusion, Summary and Future Scope")

    heading(doc, 2, "7.1 Summary of the Work")
    para(doc,
         "This project designed, implemented, and rigorously evaluated a "
         "complete deep-learning system for four-class brain tumor "
         "classification from T1-weighted MRI images. The system combines "
         "three architecturally diverse neural networks: a custom CNN "
         "trained from scratch (1.2 million parameters), an "
         "EfficientNet-B0 model adapted via two-stage transfer learning "
         "from ImageNet (4.0 million parameters), and a Swin-Tiny vision "
         "transformer fine-tuned with careful learning rate control (27.5 "
         "million parameters). These three models' softmax probability "
         "outputs are combined through a weighted-average ensemble whose "
         "optimal mixing weights [0.0, 0.9, 0.1] were determined by "
         "exhaustive grid search on the validation set.",
         align="justify")
    para(doc,
         "The system was trained and evaluated on the Brain Tumor MRI "
         "Dataset (Nickparvar, Kaggle), containing 7,200 images across "
         "four balanced classes. A stratified 80/20 split produced 4,480 "
         "training and 1,120 validation images, while 1,600 test images "
         "were held out and evaluated exactly once to produce the final "
         "reported metrics.",
         align="justify")
    para(doc,
         "Beyond model training and evaluation, the project delivered "
         "four deployment interfaces (Streamlit web app, FastAPI REST "
         "endpoint, ONNX exports, Docker container), Grad-CAM "
         "explainability for all three models, a configurable "
         "low-confidence threshold gate, and comprehensive documentation "
         "including this report, architecture diagrams, and viva "
         "preparation guides.",
         align="justify")

    heading(doc, 2, "7.2 Conclusion")
    para(doc,
         "The dominant individual model, EfficientNet-B0 with transfer "
         "learning, achieves 94.94% accuracy on the held-out test set "
         "with a macro-averaged ROC-AUC of 0.991. The three-model "
         "ensemble achieves 94.69% accuracy at 0.991 AUC, matching but "
         "not exceeding the best individual model. This result is "
         "consistent with the dominant-model problem documented in the "
         "ensemble learning literature: when one base learner "
         "substantially outperforms the others and the remaining learners "
         "make correlated errors, ensembling adds noise rather than "
         "reducing it.",
         align="justify")
    para(doc,
         "Grad-CAM analysis confirms that all three models attend to "
         "anatomically meaningful regions. EfficientNet's heatmaps are "
         "the most spatially precise; Swin's are the most diffuse. "
         "Remaining errors map to explainable failure modes, primarily "
         "the glioma-meningioma boundary where genuine clinical ambiguity "
         "exists even for human radiologists.",
         align="justify")
    para(doc,
         "The system is fully deployable through four channels: a "
         "Streamlit web interface for interactive demonstrations, a "
         "FastAPI REST backend for programmatic integration, ONNX exports "
         "for cross-platform deployment (yielding 2 to 3.5x CPU "
         "speedup), and a Docker container for cloud deployment. All four "
         "interfaces wrap the same trained checkpoints, ensuring "
         "consistent predictions regardless of the access method.",
         align="justify")
    para(doc,
         "The project demonstrates that a complete, production-quality "
         "medical imaging AI system can be built, trained, and deployed "
         "on modest hardware (a single NVIDIA GTX 1650 with 4 GB VRAM) "
         "within an academic semester, provided that careful attention "
         "is paid to engineering fundamentals: reproducible data splits, "
         "mixed-precision training, two-stage transfer learning, "
         "honest test-set evaluation, and visual explainability.",
         align="justify")

    heading(doc, 2, "7.3 Future Scope")
    para(doc,
         "While the current system delivers a robust solution, several "
         "avenues for future enhancement exist:",
         align="justify")
    bullet(doc,
           "3D Volumetric Classification: The current system classifies "
           "individual 2D slices. Future work could extend the pipeline "
           "to process full 3D MRI volumes using architectures such as "
           "3D-ResNet or V-Net, which would capture inter-slice spatial "
           "context and potentially improve classification accuracy for "
           "tumors that manifest across multiple slices.")
    bullet(doc,
           "Multi-Sequence Fusion: Clinical radiology uses multiple MRI "
           "sequences (T1, T2, FLAIR, contrast-enhanced T1) for diagnosis. "
           "Incorporating multiple sequences as input channels or through "
           "a multi-stream fusion architecture could significantly improve "
           "the system's ability to distinguish between tumor types that "
           "appear similar on T1 alone.")
    bullet(doc,
           "Self-Supervised Pretraining: Training a foundation model on a "
           "large corpus of unlabelled brain MRI data using techniques "
           "such as masked image modelling before supervised fine-tuning "
           "could yield representations better suited to the MRI domain "
           "than ImageNet-pretrained features.")
    bullet(doc,
           "Test-Time Augmentation (TTA): Averaging predictions over "
           "multiple augmented versions of the same test image (horizontal "
           "flip, small rotations) at inference time could improve "
           "accuracy at the cost of proportionally increased inference "
           "latency.")
    bullet(doc,
           "Cross-Institution Validation: The current evaluation uses a "
           "single public dataset. Validating the system on multi-centre "
           "clinical data from different scanner manufacturers and "
           "acquisition protocols would be essential before any real "
           "clinical deployment.")
    bullet(doc,
           "Federated Learning: To address privacy concerns in clinical "
           "settings, the training pipeline could be adapted for federated "
           "learning, where models are trained across multiple hospital "
           "sites without centralising patient data.")
    bullet(doc,
           "LLM-Assisted Report Generation: Integrating a large language "
           "model to generate preliminary radiology reports from the "
           "classification output and Grad-CAM findings could further "
           "assist radiologists in high-volume settings.")

    doc.add_page_break()

    # ============================================================ BIBLIOGRAPHY
    heading(doc, 1, "Bibliography")
    refs = [
        "[1] Bouhafra, S. and El Bahi, H., "
        '"Deep Learning Approaches for Brain Tumor Detection and '
        'Classification Using MRI Images (2020 to 2024): A Systematic '
        'Review," Journal of Digital Imaging Informatics in Medicine, '
        "vol. 38, pp. 1403-1433, 2025.",

        "[2] Ranjbarzadeh, R. et al., "
        '"Explainable AI and Vision Transformers for Detection and '
        'Classification of Brain Tumor: A Comprehensive Survey," '
        "Artificial Intelligence Review, Springer, 2025.",

        "[3] Babu Vimala, B., Srinivasan, S., Mathivanan, S.K. et al., "
        '"Detection and Classification of Brain Tumor Using Hybrid Deep '
        'Learning Models," Scientific Reports, vol. 13, 23029, 2023.',

        "[4] Islam, M.T. et al., "
        '"BrainNet: Precision Brain Tumor Classification with Optimized '
        'EfficientNet Architecture," International Journal of Intelligent '
        "Systems, Wiley, 2024.",

        "[5] Alnowami, M. et al., "
        '"Enhancing EfficientNetv2 with Global and Efficient Channel '
        'Attention Mechanisms for Accurate MRI-Based Brain Tumor '
        'Classification," Cluster Computing, Springer, 2024.',

        "[6] Haq, A.U. et al., "
        '"A Novel Swin Transformer Approach Utilizing Residual '
        'Multi-Layer Perceptron for Diagnosing Brain Tumors in MRI '
        'Images," International Journal of Machine Learning and '
        "Cybernetics, vol. 15, pp. 3579-3597, 2024.",

        "[7] Alsubai, S. et al., "
        '"Enhanced Magnetic Resonance Imaging-Based Brain Tumor '
        'Classification with a Hybrid Swin Transformer and ResNet50V2 '
        'Model," Applied Sciences, vol. 14, no. 22, 10154, 2024.',

        "[8] Khan, M.A. et al., "
        '"Weighted Average Ensemble Deep Learning Model for '
        'Stratification of Brain Tumor in MRI Images," Diagnostics, '
        "vol. 13, no. 7, 1320, 2023.",

        "[9] Abdella, G.M. et al., "
        '"Majority Voting Ensemble of Deep CNNs for Robust MRI-Based '
        'Brain Tumor Classification," Diagnostics, vol. 15, no. 14, '
        "1782, 2025.",

        "[10] Ali, M. et al., "
        '"Enhancing Brain Tumor Detection in MRI Images through '
        "Explainable AI Using Grad-CAM with ResNet 50,\" "
        "BMC Medical Imaging, 2024.",

        "[11] Naser, M.A. and Deen, M.J., "
        '"Explainable Deep Learning Approach for Multi-Class Brain MRI '
        'Tumor Classification and Localization Using Gradient-Weighted '
        'Class Activation Mapping," Information, vol. 14, no. 12, '
        "642, 2023.",

        "[12] Nickparvar, M., "
        '"Brain Tumor MRI Dataset," Kaggle, 2022. Available: '
        "https://www.kaggle.com/datasets/masoudnickparvar/"
        "brain-tumor-mri-dataset",

        "[13] Selvaraju, R.R. et al., "
        '"Grad-CAM: Visual Explanations from Deep Networks via '
        'Gradient-based Localization," in Proceedings of the IEEE '
        "International Conference on Computer Vision (ICCV), 2017.",

        "[14] Tan, M. and Le, Q.V., "
        '"EfficientNet: Rethinking Model Scaling for Convolutional '
        'Neural Networks," in Proceedings of the International '
        "Conference on Machine Learning (ICML), 2019.",

        "[15] Liu, Z. et al., "
        '"Swin Transformer: Hierarchical Vision Transformer Using '
        'Shifted Windows," in Proceedings of the IEEE International '
        "Conference on Computer Vision (ICCV), 2021.",

        "[16] Loshchilov, I. and Hutter, F., "
        '"Decoupled Weight Decay Regularization," in Proceedings of '
        "the International Conference on Learning Representations "
        "(ICLR), 2019.",

        "[17] International Association of Cancer Registries (IARC) and "
        "Indian Journal of Neurology, 2025; BW Healthcare World, "
        '"Battling the Brain Tumor Burden: Challenges and Advancements '
        'in India," 2025.',

        "[18] Micikevicius, P. et al., "
        '"Mixed Precision Training," in Proceedings of the '
        "International Conference on Learning Representations "
        "(ICLR), 2018.",

        "[19] Ahmed, S. et al., "
        '"Advancing Brain Tumor MRI Classification using SwRD: '
        'A Parallel Swin Transformer-ResNet Approach," Virtual Reality '
        "and Intelligent Hardware, vol. 7, no. 5, pp. 501-522, 2025.",

        "[20] Global Cancer Observatory (GLOBOCAN), "
        '"Cancer Today: Estimated Number of New Cases in 2020, '
        'Worldwide," International Agency for Research on Cancer, '
        "World Health Organization, 2022.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(10)

    # ============================================================ SAVE
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size:  {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
