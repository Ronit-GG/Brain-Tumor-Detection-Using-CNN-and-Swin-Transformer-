"""
Convert PROJECT_REPORT_humanized.docx -> PROJECT_REPORT_humanized.tex

Design goals:
  * Preserve the HUMANIZED wording verbatim (structure, headings, tables).
  * Fix ONLY the specific errors that were flagged, nothing else:
      1. Strip invisible / zero-width characters (U+2063 etc.)
      2. en-dash / em-dash  -> hyphen
      3. curly quotes        -> straight quotes
      4. Spanish heading "Diagramas de flujo de datos" -> "Data Flow Diagrams"
      5. Garbled "images.ics.edict REST endpoint" -> "images. Expose a /predict REST endpoint"
      6. Factual error: "validation accuracy of 1.07% ... test accuracy of 5.06%"
                        -> "validation error of 1.07% ... test error of 5.06%"
      7. Corrupted AdamW hyperparameter line -> restore beta1/beta2/weight-decay
      8. Broken heading "Software Requirement Specification 5.2.1 SRS"
                        -> "Software Requirements Specification (SRS)"
      9. Typo "Brain Tumor MRI Datas" -> "Brain Tumor MRI Dataset"
     10. Section renames flagged: "Functional Needs"->"Functional Requirements",
         "Need for Interfaces"->"Interface Requirements", "In Summary"->"Conclusion"
  * Figures are inserted using the COMPRESSED images (outputs/report_figures_compressed)
    so the document compiles inside the free Overleaf 60-second limit.

Math fragments in the source doc are inconsistent (some wrapped in $...$, some
not, one unbalanced $). To guarantee compilation they are flattened to readable
upright text (e.g. "p_ens = w_cnn x p_cnn + ..."). Wording is unchanged.

Usage:
    cd "C:\\Brain Tumor Detection"
    .\.venv\Scripts\python.exe docs\humanized_to_latex.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = Path(__file__).resolve().parent / "PROJECT_REPORT_humanized.docx"
OUT = Path(__file__).resolve().parent / "PROJECT_REPORT_humanized.tex"


# --------------------------------------------------------------------------- #
# Figure placement: maps a section number -> list of (basename, caption, width)
# Filenames refer to the COMPRESSED bundle (outputs/report_figures_compressed/).
# Photo figures were saved as .jpg by the compressor; diagrams stayed .png.
# --------------------------------------------------------------------------- #
JPG = {"02", "04", "06", "07", "28", "29"}

def figfile(num_prefix: str, slug: str) -> str:
    ext = "jpg" if num_prefix in JPG else "png"
    return f"Figure_{num_prefix}_{slug}.{ext}"

FIG_MAP = {
    "1.1":   [(figfile("08", "system_flow_diagram"),
               "High-level system flow: MRI input, preprocessing, three parallel "
               "models, weighted ensemble, predicted class with confidence and "
               "Grad-CAM heatmaps.", 0.92)],
    "3.3":   [(figfile("02", "sample_mri_per_class"),
               "Representative MRI slices, one per class.", 0.85)],
    "3.5":   [(figfile("01", "class_distribution"),
               "Per-class image counts across the Training, Validation and Test "
               "splits.", 0.8)],
    "3.6":   [(figfile("03", "image_size_distribution"),
               "Distribution of original image dimensions.", 0.8),
              (figfile("05", "pixel_intensity_histograms"),
               "Per-class pixel-intensity histograms showing substantial overlap.",
               0.8)],
    "3.8":   [(figfile("06", "preprocessing_effect"),
               "Visual effect of the deterministic preprocessing pipeline.", 0.8)],
    "3.9":   [(figfile("07", "augmentation_examples"),
               "Examples of training-time data augmentation.", 0.85)],
    "4.7":   [(figfile("31", "gantt_chart"),
               "Project schedule across the semester.", 0.92)],
    "5.3.1": [(figfile("09", "architecture_custom_cnn"),
               "Custom CNN architecture.", 0.92),
              (figfile("10", "architecture_efficientnet_b0"),
               "EfficientNet-B0 transfer-learning architecture.", 0.92),
              (figfile("11", "architecture_swin_tiny"),
               "Swin-Tiny architecture.", 0.92)],
    "5.3.3": [(figfile("27", "ensemble_weight_heatmap"),
               "Validation accuracy across the ensemble-weight grid search.", 0.7)],
    "5.3.4": [(figfile("28", "gradcam_correct"),
               "Grad-CAM overlays for correctly classified samples (one per class).",
               0.9)],
    "5.4.1": [(figfile("15", "usecase_radiologist"),
               "Use-case diagram: Radiologist / Clinician role.", 0.8),
              (figfile("16", "usecase_researcher"),
               "Use-case diagram: Researcher / Developer role.", 0.8),
              (figfile("17", "usecase_admin"),
               "Use-case diagram: Administrator role.", 0.8)],
    "5.4.2": [(figfile("18", "dfd_training"),
               "Data Flow Diagram: model training pipeline.", 0.92),
              (figfile("30", "dfd_inference"),
               "Data Flow Diagram: single-image inference.", 0.92)],
    "6.2":   [(figfile("12", "cnn_training_curves"),
               "Custom CNN training and validation curves.", 0.8),
              (figfile("13", "efficientnet_training_curves"),
               "EfficientNet-B0 training curves across both stages.", 0.8),
              (figfile("14", "swin_training_curves"),
               "Swin-Tiny training curves across both stages.", 0.8)],
    "6.4":   [(figfile("19", "confusion_matrix_cnn"),
               "Test-set confusion matrix: Custom CNN.", 0.6),
              (figfile("20", "confusion_matrix_efficientnet"),
               "Test-set confusion matrix: EfficientNet-B0.", 0.6),
              (figfile("21", "confusion_matrix_swin"),
               "Test-set confusion matrix: Swin-Tiny.", 0.6),
              (figfile("22", "confusion_matrix_ensemble"),
               "Test-set confusion matrix: Weighted ensemble.", 0.6),
              (figfile("24", "per_class_f1_bars"),
               "Per-class F1 scores across all models.", 0.8)],
    "6.5":   [(figfile("29", "gradcam_misclassified"),
               "Grad-CAM overlays for misclassified samples, illustrating failure "
               "modes.", 0.9)],
    "6.6":   [(figfile("23", "roc_curves_test"),
               "ROC curves on the test set for all models and the ensemble.", 0.75),
              (figfile("25", "final_comparison_bars"),
               "Final comparison of test-set accuracy, F1 and macro AUC.", 0.8),
              (figfile("26", "efficiency_pareto"),
               "Accuracy versus inference-latency trade-off.", 0.75)],
}


# --------------------------------------------------------------------------- #
# Text cleaning + the flagged fixes
# --------------------------------------------------------------------------- #
INVISIBLE = ["\u200b", "\u200c", "\u200d", "\u2060", "\u2061", "\u2062",
             "\u2063", "\u2064", "\ufeff"]

MATH_CMDS = [
    ("\\mathbb{R}", "R"), ("\\mathbb", "R"),
    ("\\times", "x"), ("\\cdot", "*"),
    ("\\geq", ">="), ("\\leq", "<="), ("\\ge", ">="), ("\\le", "<="),
    ("\\in", " in "), ("\\sum", "sum"), ("\\frac", ""),
    ("\\,", " "), ("\\;", " "),
]


def normalise(text: str) -> str:
    """Steps 1-2: strip invisible chars, normalise dashes and curly quotes."""
    for ch in INVISIBLE:
        text = text.replace(ch, "")
    text = text.replace("\u2013", "-").replace("\u2014", "-")   # en / em dash
    text = text.replace("\u2018", "'").replace("\u2019", "'")   # single curly
    text = text.replace("\u201c", '"').replace("\u201d", '"')   # double curly
    text = text.replace("\u2022", "")                            # stray bullet glyph
    return text


def apply_fixes(text: str) -> str:
    """Steps 4-10: the specific flagged corrections (operate on clean text)."""
    # 4. Spanish heading
    text = text.replace("Diagramas de flujo de datos", "Data Flow Diagrams")
    # 5. Garbled REST sentence
    text = text.replace("images.ics.edict REST endpoint",
                        "images. Expose a /predict REST endpoint")
    # 6. Factual error: error rates, not accuracies
    text = text.replace("validation accuracy of 1.07%", "validation error of 1.07%")
    text = text.replace("test accuracy of 5.06%", "test error of 5.06%")
    # 7. Restore the corrupted AdamW hyperparameter sentence
    text = re.sub(
        r"The main hyperparameters are:.*?The learning rate schedule",
        "The main hyperparameters are: beta_1 = 0.9, beta_2 = 0.999, and "
        "weight decay = 1e-4. The learning rate schedule",
        text, flags=re.DOTALL)
    # 8. Broken SRS heading
    text = text.replace("Software Requirement Specification 5.2.1 SRS",
                        "Software Requirements Specification (SRS)")
    # 9. Typo
    text = text.replace("Brain Tumor MRI Datas ", "Brain Tumor MRI Dataset ")
    # 10. Section renames that were flagged
    text = text.replace("Functional Needs", "Functional Requirements")
    text = text.replace("Need for Interfaces", "Interface Requirements")
    text = text.replace("In Summary", "Conclusion")
    return text


def flatten_math(text: str) -> str:
    """Make inconsistent inline-math fragments compile as plain upright text."""
    text = text.replace("$", "")
    for a, b in MATH_CMDS:
        text = text.replace(a, b)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\\[a-zA-Z]+", "", text)   # drop any leftover \command
    text = text.replace("\\", "")             # drop any stray backslash
    return text


def latex_escape(text: str) -> str:
    """Escape the LaTeX specials that may remain in ordinary prose."""
    text = text.replace("&", r"\&").replace("%", r"\%").replace("#", r"\#")
    text = text.replace("_", r"\_")
    text = text.replace("~", r"\textasciitilde{}")
    text = text.replace("^", r"\textasciicircum{}")
    return text


def clean(text: str) -> str:
    return latex_escape(flatten_math(apply_fixes(normalise(text))))


# --------------------------------------------------------------------------- #
# Document walking
# --------------------------------------------------------------------------- #
def iter_blocks(doc: _Doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def heading_level(p: Paragraph) -> int:
    name = (p.style.name or "").lower()
    if name.startswith("heading 1") or name == "title":
        return 1
    if name.startswith("heading 2"):
        return 2
    if name.startswith("heading 3"):
        return 3
    return 0


def is_bullet(p: Paragraph) -> bool:
    name = (p.style.name or "").lower()
    if "list" in name:
        return True
    return p.text.lstrip().startswith("\u2022")


def section_number(heading_text: str) -> str | None:
    m = re.match(r"\s*(\d+(?:\.\d+)*)", heading_text)
    return m.group(1) if m else None


# --------------------------------------------------------------------------- #
# Emitters
# --------------------------------------------------------------------------- #
PREAMBLE = r"""\documentclass[12pt,a4paper,oneside]{report}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{lmodern}
\usepackage[margin=1in]{geometry}
\usepackage{setspace}
\usepackage{graphicx}
\usepackage{float}
\usepackage{caption}
\usepackage{booktabs}
\usepackage{tabularx}
\usepackage{array}
\usepackage{enumitem}
\usepackage{fancyhdr}
\usepackage[hidelinks]{hyperref}
\usepackage{url}

\onehalfspacing
\graphicspath{{figures/}}
\setlength{\parskip}{6pt}
\setlength{\parindent}{0pt}
\sloppy

\newcolumntype{Y}{>{\raggedright\arraybackslash}X}

\pagestyle{fancy}
\fancyhf{}
\fancyhead[L]{\small Brain Tumor Detection}
\fancyhead[R]{\small \thepage}
\renewcommand{\headrulewidth}{0.4pt}

\captionsetup{font=small,labelfont=bf}

\begin{document}
"""


class Builder:
    def __init__(self) -> None:
        self.out: list[str] = []
        self.in_itemize = False
        self.title_done = False
        self.title_lines: list[str] = []
        self.seen_first_heading = False

    def close_itemize(self) -> None:
        if self.in_itemize:
            self.out.append(r"\end{itemize}")
            self.in_itemize = False

    def open_itemize(self) -> None:
        if not self.in_itemize:
            self.out.append(r"\begin{itemize}[leftmargin=*]")
            self.in_itemize = True

    def emit_titlepage(self) -> None:
        self.out.append(r"\begin{titlepage}")
        self.out.append(r"\centering")
        self.out.append(r"\vspace*{1.5cm}")
        for i, raw in enumerate(self.title_lines):
            if not raw.strip():
                continue
            t = clean(raw)
            if i == 0:
                self.out.append(r"{\LARGE\bfseries " + t + r"\par}")
                self.out.append(r"\vspace{1.2cm}")
            else:
                self.out.append(t + r"\\[4pt]")
        self.out.append(r"\end{titlepage}")
        self.out.append(r"\pagenumbering{roman}")
        self.out.append(r"\tableofcontents")
        self.out.append(r"\listoffigures")
        self.out.append(r"\clearpage")
        self.out.append(r"\pagenumbering{arabic}")
        self.title_done = True

    def emit_figures_for(self, num: str) -> None:
        for fname, cap, width in FIG_MAP.get(num, []):
            self.out.append(r"\begin{figure}[H]")
            self.out.append(r"  \centering")
            self.out.append(rf"  \includegraphics[width={width}\textwidth]{{{fname}}}")
            self.out.append(rf"  \caption{{{cap}}}")
            self.out.append(r"\end{figure}")

    def add_heading(self, p: Paragraph, level: int) -> None:
        self.close_itemize()
        if not self.seen_first_heading:
            # Everything collected so far is the title page.
            self.seen_first_heading = True
            if not self.title_done:
                self.emit_titlepage()
        raw = p.text
        # Heading-targeted fixes are inside apply_fixes (called by clean()).
        text = clean(raw)
        if level == 1:
            self.out.append(r"\chapter*{" + text + "}")
            self.out.append(r"\addcontentsline{toc}{chapter}{" + text + "}")
        elif level == 2:
            self.out.append(r"\section*{" + text + "}")
            self.out.append(r"\addcontentsline{toc}{section}{" + text + "}")
        else:
            self.out.append(r"\subsection*{" + text + "}")
            self.out.append(r"\addcontentsline{toc}{subsection}{" + text + "}")
        num = section_number(apply_fixes(normalise(raw)))
        if num:
            self.emit_figures_for(num)

    def add_paragraph(self, p: Paragraph) -> None:
        raw = p.text
        if not raw.strip():
            return
        if not self.seen_first_heading:
            self.title_lines.append(raw)
            return
        if is_bullet(p):
            self.open_itemize()
            item = clean(raw.lstrip().lstrip("\u2022").strip())
            self.out.append(r"  \item " + item)
            return
        self.close_itemize()
        stripped = raw.strip()
        if stripped.lower().startswith("http"):
            self.out.append(r"\url{" + stripped + "}")
        else:
            self.out.append(clean(raw))

    def add_table(self, t: Table) -> None:
        self.close_itemize()
        ncols = len(t.columns)
        size = r"\footnotesize" if ncols >= 5 else r"\small"
        colspec = "Y" * ncols
        self.out.append(r"\begin{center}")
        self.out.append(size)
        self.out.append(r"\begin{tabularx}{\textwidth}{" + colspec + "}")
        self.out.append(r"\toprule")
        for ri, row in enumerate(t.rows):
            cells = [clean(c.text.replace("\n", " ").strip()) for c in row.cells]
            if ri == 0:
                cells = [r"\textbf{" + c + "}" for c in cells]
            self.out.append(" & ".join(cells) + r" \\")
            if ri == 0:
                self.out.append(r"\midrule")
        self.out.append(r"\bottomrule")
        self.out.append(r"\end{tabularx}")
        self.out.append(r"\end{center}")

    def build(self, doc: _Doc) -> str:
        for block in iter_blocks(doc):
            if isinstance(block, Paragraph):
                lvl = heading_level(block)
                if lvl:
                    self.add_heading(block, lvl)
                else:
                    self.add_paragraph(block)
            elif isinstance(block, Table):
                self.add_table(block)
        self.close_itemize()
        return PREAMBLE + "\n".join(self.out) + "\n\\end{document}\n"


def main() -> None:
    doc = Document(str(SRC))
    tex = Builder().build(doc)
    OUT.write_text(tex, encoding="utf-8")
    print(f"Saved: {OUT}")
    print(f"Size:  {OUT.stat().st_size/1024:.0f} KB")
    # Sanity checks
    bad = {
        "dollar signs ($)": tex.count("$"),
        "invisible U+2063": tex.count("\u2063"),
        "Spanish heading":  tex.count("Diagramas"),
        "ics.edict garble": tex.count("ics.edict"),
        "en-dash":          tex.count("\u2013"),
        "curly quotes":     sum(tex.count(c) for c in "\u2018\u2019\u201c\u201d"),
    }
    print("Sanity checks (all should be 0):")
    for k, v in bad.items():
        print(f"  {k:20s}: {v}")
    print("Fix confirmations (should be >=1):")
    print(f"  'validation error of 1.07%' : {tex.count('validation error of 1.07')}")
    print(f"  'Data Flow Diagrams'        : {tex.count('Data Flow Diagrams')}")
    print(f"  'Expose a /predict'         : {tex.count('Expose a /predict')}")
    print(f"  'beta_1 = 0.9' (escaped)    : {tex.count('beta')}")
    print(f"  brace balance {{ vs }}        : {tex.count('{')} vs {tex.count('}')}")


if __name__ == "__main__":
    main()
