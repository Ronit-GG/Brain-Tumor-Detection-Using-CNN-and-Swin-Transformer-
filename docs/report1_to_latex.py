"""
Convert docs/PROJECT_REPORT 1.docx -> docs/PROJECT_REPORT_final.tex

- Front matter (title through acknowledgement): user-provided LaTeX template.
- Body: from Abstract onward, wording from the edited docx.
- Only necessary fixes (same class as humanized_to_latex.py + report-1 extras).
- Figures: compressed bundle (upload outputs/report_figures_compressed/ to Overleaf figures/).

Usage:
    cd "C:\\Brain Tumor Detection"
    .\\.venv\\Scripts\\python.exe docs\\report1_to_latex.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from humanized_to_latex import (
    FIG_MAP,
    Builder,
    PREAMBLE,
    apply_fixes,
    clean,
    figfile,
    heading_level,
    is_bullet,
    iter_blocks,
    normalise,
    section_number,
)

SRC = Path(__file__).resolve().parent / "PROJECT_REPORT 1.docx"
OUT = Path(__file__).resolve().parent / "PROJECT_REPORT_final.tex"

PROJECT_TITLE = (
    "Brain Tumor Detection from MRI Images using an Ensemble of CNN, "
    "Transfer Learning and Swin Transformer Models"
)

# User template through acknowledgement (project title + real acknowledgement text).
FRONT_MATTER = r"""
% --------------------------------------------------
% PAGE 1 : TITLE PAGE
% --------------------------------------------------

\begin{titlepage}
\centering

\vspace*{1cm}

{\Huge \textbf{""" + PROJECT_TITLE.replace("_", r"\_") + r"""}}\\[1cm]

{\Large 7th/8th Semester Project Report}\\[1cm]

submitted to\\[0.3cm]

{\Large \textbf{Techno India University, West Bengal}}\\[1cm]

In partial fulfilment of the requirements\\
for the award of the degree of\\[0.3cm]

{\Large \textbf{Bachelor of Technology}}\\[0.3cm]

in\\[0.3cm]

{\Large \textbf{Computer Science \& Engineering}}\\[1cm]

by\\[0.5cm]

\begin{tabular}{c}
\textbf{Ronit Singh (University Roll No. - 221001001297)} \\[0.3cm]
\textbf{Ahana Sarkar (University Roll No. - 221001001222)} \\[0.3cm]
\textbf{Swikriti Roy (University Roll No. - 221001001279)} \\[0.3cm]
\textbf{Dibyashree Dey (University Roll No. - 221001001346)} \\[0.3cm]
\textbf{Shreyansh Kumar Gupta (University Roll No. - 221001001290)}
\end{tabular}

\vspace{1.5cm}

Under the guidance of\\[0.3cm]

{\Large \textbf{Dr. Priyanka Saha}}\\[1.5cm]

Department of Computer Science \& Engineering\\
\textbf{TECHNO INDIA UNIVERSITY, WEST BENGAL}\\[1cm]

January 2026\\[0.5cm]

\copyright\ 2026, Techno India University. All rights reserved.

\end{titlepage}

% --------------------------------------------------
% PAGE 2 : BLANK PAGE
% --------------------------------------------------

\newpage
\thispagestyle{empty}
\mbox{}

% --------------------------------------------------
% PAGE 3 : DEDICATION
% --------------------------------------------------

\newpage
\thispagestyle{empty}

\vspace*{7cm}

\begin{center}
{\Huge Dedicated to}\\[1cm]
Our parents, families, and teachers who supported us throughout this journey.
\end{center}

% --------------------------------------------------
% PAGE 4 : BLANK PAGE
% --------------------------------------------------

\newpage
\thispagestyle{empty}
\mbox{}

% --------------------------------------------------
% PAGE 5 : DECLARATION OF AUTHORSHIP
% --------------------------------------------------

\newpage

\begin{center}
{\Large \textbf{DECLARATION OF AUTHORSHIP}}
\end{center}

\vspace{1cm}

We hereby declare that the project report entitled ``""" + PROJECT_TITLE + r"""'' is an authentic
record of our own work carried out at the Department of Computer Science and
Engineering, Techno India University, West Bengal, during the 8th semester
of the academic year 2025--2026 under the supervision of Dr. Priyanka Saha,
Assistant Professor.

\vspace{0.5cm}

We further declare that the matter embodied in this project report has not been submitted
by us for the award of any other degree or diploma of this or any other Institute/University.
All the information have been obtained and presented in accordance with academic
rules and ethical conduct. We also declare that, as required by these rules and conduct,
we have fully cited and referenced all materials and results that are not original to this
work. The dataset used in this project is publicly available under its original license; no
patient-identifiable data was collected, generated, or stored.

\vspace{1cm}

\textbf{Place:} EM-4/1, Sector V, Bidhannagar, Kolkata, West Bengal -- 700091

\vspace{0.5cm}

\textbf{Date:} January 2026

\vspace{1cm}

\textbf{Signatures:}

\vspace{1cm}

\begin{tabular}{p{7cm}p{7cm}}
Signature: & Signature: \\
Name: Ronit Singh & Name: Ahana Sarkar \\
Roll No: 221001001297 & Roll No: 221001001222
\end{tabular}

\vspace{1cm}

\begin{tabular}{p{7cm}p{7cm}}
Signature: & Signature: \\
Name: Swikriti Roy & Name: Dibyashree Dey \\
Roll No: 221001001279 & Roll No: 221001001346
\end{tabular}

\vspace{1cm}

\begin{tabular}{p{7cm}}
Signature: \\
Name: Shreyansh Kumar Gupta \\
Roll No: 221001001290
\end{tabular}

% --------------------------------------------------
% PAGE 6 : BLANK PAGE
% --------------------------------------------------

\newpage
\thispagestyle{empty}
\mbox{}

% --------------------------------------------------
% PAGE 7 : CERTIFICATE OF RECOMMENDATION
% --------------------------------------------------

\newpage

\begin{center}
{\Large \textbf{CERTIFICATE OF RECOMMENDATION}}
\end{center}

\vspace{1cm}

This is to certify that the work embodied in this thesis entitled ``""" + PROJECT_TITLE + r"""''
has been satisfactorily completed by Ronit Singh (Roll No. 221001001297),
Ahana Sarkar (Roll No. 221001001222),
Swikriti Roy (Roll No. 221001001279),
Dibyashree Dey (Roll No. 221001001346),
and Shreyansh Kumar Gupta (Roll No. 221001001290).

It is a bonafide piece of work
carried out under my supervision and guidance at Techno India University, Kolkata,
for partial fulfilment of the requirements for the awarding of the Bachelor of Technology
in Computer Science \& Engineering degree of the Department of
Computer Science and Engineering, Techno India University, during the academic year
2025--2026.

\vspace{2cm}

\textbf{Dr. Priyanka Saha}\\
Assistant Professor, Department of Computer Science and Engineering,\\
Techno India University, Kolkata, West Bengal, India.\\
(Supervisor)

\vspace{2cm}

\textbf{Forwarded By:}

\vspace{1cm}

\textbf{[Name of Head of Department]}\\
HoD, Department of Computer Science and Engineering,\\
Techno India University, Kolkata, West Bengal, India.

% --------------------------------------------------
% PAGE 8 : BLANK PAGE
% --------------------------------------------------

\newpage
\thispagestyle{empty}
\mbox{}

% --------------------------------------------------
% PAGE 9 : ACKNOWLEDGEMENT
% --------------------------------------------------

\newpage

\begin{center}
{\Large \textbf{ACKNOWLEDGEMENT}}
\end{center}

\vspace{1cm}

We would like to first express our sincere gratitude to our project supervisor,
Dr. Priyanka Saha, Assistant Professor, Department of Computer Science and Engineering,
Techno India University, West Bengal. Their continued guidance, technical insight, and
constructive feedback during every stage of this project were instrumental in shaping the
final outcome. Their willingness to discuss difficult engineering decisions, including the
diagnosis of subtle numerical-stability issues during mixed-precision training, helped us
approach the work with both rigour and curiosity.

\vspace{0.5cm}

We take this opportunity to express our gratitude to all faculty members of the Department
of Computer Science and Engineering for their support, the encouragement they provided
during our coursework, and the foundation they laid in machine learning, computer vision,
and software engineering.

\vspace{0.5cm}

We acknowledge the open-source community whose tools made this work feasible: the PyTorch
and timm library maintainers, the authors of pytorch-grad-cam, the Streamlit and FastAPI
teams, and Masoud Nickparvar for releasing the Brain Tumor MRI Dataset on Kaggle under a
permissive license that supports academic research.

\vspace{0.5cm}

Finally, we thank our parents and families for their unceasing encouragement, patience,
and support throughout the semester.

\newpage
\thispagestyle{empty}
\mbox{}

\pagenumbering{roman}
"""


def apply_report1_fixes(text: str) -> str:
    """Extra fixes found only in PROJECT_REPORT 1.docx."""
    text = apply_fixes(text)
    text = text.replace("Needs Functionnelles", "Functional Requirements")
    text = text.replace("Non Functional Characteristics", "Non-Functional Requirements")
    text = text.replace("Model of the software.", "Model of the Software")
    text = text.replace(
        "(2024): BrainNet Optimised EfficientNet Architecture.",
        "Islam et al. (2024): BrainNet Optimised EfficientNet Architecture.",
    )
    if text.strip() == "Data flow diagrams":
        text = "Data Flow Diagrams"
    # Garbled REST line (bullet may be stripped before this runs)
    text = text.replace(
        "ics.edict REST endpoint that accepts",
        "Expose a /predict REST endpoint that accepts",
    )
    text = re.sub(r"Brain Tumou?r MRI Datas\b", "Brain Tumor MRI Dataset", text)
    text = text.replace("Datasetet", "Dataset")
    return text


def clean_report1(text: str) -> str:
    from humanized_to_latex import flatten_math, latex_escape

    return latex_escape(flatten_math(apply_report1_fixes(normalise(text))))


class Report1Builder(Builder):
    """Skip docx front matter; do not auto-generate title page."""

    def __init__(self) -> None:
        super().__init__()
        self.phase = "skip"  # skip | abstract | body
        self.abstract_paras: list[str] = []

    def add_heading(self, p: Paragraph, level: int) -> None:
        raw = p.text.strip()
        low = raw.lower()
        if self.phase == "skip":
            if low == "abstract":
                self.phase = "abstract"
                self.seen_first_heading = True
                self.title_done = True
            return
        if self.phase == "abstract":
            if low.startswith("chapter"):
                self.phase = "body"
            else:
                return
        self.close_itemize()
        text = clean_report1(raw)
        if level == 1:
            if text.lower().startswith("chapter"):
                self.out.append(r"\chapter{" + text + "}")
            else:
                self.out.append(r"\chapter*{" + text + "}")
                self.out.append(r"\addcontentsline{toc}{chapter}{" + text + "}")
        elif level == 2:
            self.out.append(r"\section{" + text + "}")
        else:
            self.out.append(r"\subsection{" + text + "}")
        num = section_number(apply_report1_fixes(normalise(raw)))
        if num:
            self.emit_figures_for(num)

    def add_paragraph(self, p: Paragraph) -> None:
        raw = p.text
        if not raw.strip():
            return
        if self.phase == "skip":
            return
        if self.phase == "abstract":
            self.abstract_paras.append(raw)
            return
        if is_bullet(p):
            self.open_itemize()
            item = clean_report1(raw.lstrip().lstrip("\u2022").strip())
            self.out.append(r"  \item " + item)
            return
        self.close_itemize()
        stripped = raw.strip()
        if stripped.lower().startswith("http"):
            self.out.append(r"\url{" + stripped + "}")
        else:
            self.out.append(clean_report1(raw))

    def add_table(self, t: Table) -> None:
        if self.phase != "body":
            return
        self.close_itemize()
        ncols = len(t.columns)
        size = r"\footnotesize" if ncols >= 5 else r"\small"
        colspec = "Y" * ncols
        self.out.append(r"\begin{center}")
        self.out.append(size)
        self.out.append(r"\begin{tabularx}{\textwidth}{" + colspec + "}")
        self.out.append(r"\toprule")
        for ri, row in enumerate(t.rows):
            cells = [
                clean_report1(c.text.replace("\n", " ").strip()) for c in row.cells
            ]
            if ri == 0:
                cells = [r"\textbf{" + c + "}" for c in cells]
            self.out.append(" & ".join(cells) + r" \\")
            if ri == 0:
                self.out.append(r"\midrule")
        self.out.append(r"\bottomrule")
        self.out.append(r"\end{tabularx}")
        self.out.append(r"\end{center}")

    def abstract_block(self) -> list[str]:
        lines = [
            r"\chapter*{Abstract}",
            r"\addcontentsline{toc}{chapter}{Abstract}",
        ]
        for raw in self.abstract_paras:
            if raw.strip():
                lines.append(clean_report1(raw))
        lines += [
            r"\tableofcontents",
            r"\listoffigures",
            r"\listoftables",
            r"\clearpage",
            r"\pagenumbering{arabic}",
        ]
        return lines

    def build(self, doc: _Doc) -> str:
        for block in iter_blocks(doc):
            if isinstance(block, Paragraph):
                lvl = heading_level(block)
                if lvl:
                    self.add_heading(block, lvl)
                else:
                    self.add_paragraph(block)
            elif isinstance(block, Table):
                self.add_table(block)
        self.close_itemize()
        parts = [PREAMBLE, FRONT_MATTER]
        parts.extend(self.abstract_block())
        parts.extend(self.out)
        return "\n".join(parts) + "\n\\end{document}\n"


def main() -> None:
    doc = Document(str(SRC))
    tex = Report1Builder().build(doc)
    OUT.write_text(tex, encoding="utf-8")
    print(f"Saved: {OUT}")
    print(f"Size:  {OUT.stat().st_size / 1024:.0f} KB")
    bad = {
        "dollar signs": tex.count("$"),
        "U+2063": tex.count("\u2063"),
        "Diagramas": tex.count("Diagramas"),
        "ics.edict": tex.count("ics.edict"),
        "Needs Functionnelles": tex.count("Needs Functionnelles"),
        "MRI Datas typo": tex.count("MRI Datas"),
        "Datasetet": tex.count("Datasetet"),
        "validation accuracy of 1.07": tex.count("validation accuracy of 1.07"),
    }
    print("Sanity (should be 0):")
    for k, v in bad.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
